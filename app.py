import os
import re
import time
import datetime
import json
import shutil
from io import BytesIO
from dotenv import load_dotenv

# Third-party libraries
from flask import Flask, render_template, request, jsonify, send_from_directory
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Length, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from PIL import Image
import openai
import requests

# --- Global Configuration and Flask App Initialization ---

app = Flask(__name__)
load_dotenv()
import google.generativeai as genai

# Create an uploads folder if it doesn't exist
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Global tags
DOCUMENTATION_TAGS = [
    'doc_introduction', 'doc_objective', 'doc_scope', 'doc_techstack',
    'doc_feasibility', 'doc_system_features', 'doc_modules', 'doc_usecase',
    'doc_advantage', 'doc_hardware_req', 'doc_software_req'
]
CODE_LIST_TAG = "[[CODE_IMPLEMENTATION_BLOCK]]"
SCREENSHOT_BLOCK_TAG = "[[SCREENSHOT_BLOCK]]" # New global tag for screenshot block

# --- Configuration Loading Function ---
def load_template_config():
    """Loads all template configurations from templates.json."""
    config_path = os.path.join(os.path.dirname(__file__), 'templates.json')
    try:
        with open(config_path, 'r') as f:
            return json.load(f)
    except FileNotFoundError:
        print(f"ERROR: templates.json not found at {config_path}. Cannot load templates.")
        return []
    except json.JSONDecodeError:
        print("ERROR: Failed to decode templates.json. Check JSON syntax.")
        return []

# Load configurations once at startup
TEMPLATE_CONFIGS = load_template_config()


# --- Helper Functions (MOCKED) ---

def get_ai_diagram_url(project_name, diagram_type):
    """MOCK: This function always returns None to ensure the local placeholder image is used."""
    return None

def download_and_save_image(url, filename, upload_folder):
    """Saves a local placeholder image, regardless of the 'url' provided."""
    target_path = os.path.join(upload_folder, filename)

    dummy_image_path = os.path.join(os.path.dirname(__file__), 'placeholder.png')

    if not os.path.exists(dummy_image_path):
        try:
            Image.new('RGB', (600, 400), color = 'red').save(dummy_image_path, 'PNG')
        except Exception:
            return False

    shutil.copyfile(dummy_image_path, target_path)
    return True

def save_placeholder_image(filename):
    return download_and_save_image(url=None, filename=filename, upload_folder=app.config['UPLOAD_FOLDER'])

# --- AI HELPER FUNCTION (GOOGLE API) ---

def get_ai_generated_documentation(project_title):
    """
    Calls the Google Generative AI (Gemini) API to get documentation.
    """
    # 1. Configure the API key
    try:
        api_key = os.getenv("GOOGLE_API_KEY")
        if not api_key:
            return {"error": "GOOGLE_API_KEY is not set. Please check your .env file."}
        genai.configure(api_key=api_key)
    except Exception as e:
        return {"error": f"API configuration failed: {e}"}

    # 2. Use updated model name + JSON output config
    model = genai.GenerativeModel(
    'gemini-2.0-flash',  # <-- Use this one
    generation_config={"response_mime_type": "application/json"}
)

    # 3. Create the prompt
    json_keys = DOCUMENTATION_TAGS 

    
    system_prompt = f"""
You are an expert technical writer for student software projects.
You will be given a project title.

Your task is to generate complete project documentation in a SINGLE JSON object.

The JSON object MUST contain EXACTLY these keys: {json_keys}

Each key must contain a well-written, structured section that follows the minimum word count rules:

- introduction: at least 180 words
- objective: at least 130 words
- scope: at least 180 words
- technology_stack: at least 140 words
- feasibility_study: at least 250 words
- system_features: at least 150 words
- modules: at least 300 words
- usecase: at least 120 words
- advantages: at least 250 words
- hardware_and_software_requirement: at least 250 words

----------------------------------------------------------
### OBJECTIVE FORMAT
he **objective** section MUST strictly follow this format:

Objective
The main objectives of the <PROJECT_TITLE> are to:
• Provide a bullet-point list of objectives.
• Each objective should begin with a capital letter.
• Each point must be written in full sentences.
• The list should contain at least 6-8 bullet points.
• Tone must match academic project documentation.
• The overall word count must still be at least 130 words.

Use the exact heading “Objective” at the top of the section, followed by the description and bullet points, similar to the style shown:

Objective
The main objectives of the News Portal Website with Editor Dashboard (The Daily Chronicle) are to:
Provide real-time news...
Enable editors to contribute...
(continue bullet points...)

----------------------------------------------------------
### SCOPE FORMAT
The **scope** section MUST strictly follow this structure and tone:

The <PROJECT_TITLE> aims to provide a comprehensive, interactive, and scalable platform for its intended purpose. Its scope includes:
• A bullet-style list of features, categories, or capabilities included in the project.
• Each scope item must begin with a label (e.g., “News Categories:”), followed by a descriptive sentence.
• The writing style MUST resemble this example:

The Daily Chronicle aims to provide a comprehensive, interactive, and scalable platform for delivering real-time news while enabling editors to publish their own content. Its scope includes:
News Categories: Users can browse news across categories such as Technology, Business, Sports, Entertainment, and Health.
Recently Added Section: Editors can contribute their own articles, stored in Firebase Firestore.
Date-Based Filters: Users can filter articles by publication date.
Search Functionality: Users can perform keyword-based searches.
Responsive User Interface: Optimized for desktops, tablets, and mobile devices.
Cross-Platform Availability: Accessible via any modern browser.
Scalability and Enhancements: Expandable to include authentication, comments, notifications, etc.

• The section must remain at least 180 words.
• Maintain academic project-documentation tone.

----------------------------------------------------------
### TECHNOLOGY STACK FORMAT
The **technology_stack** section MUST follow this exact structure:

The project is implemented using the following technologies:

Frontend (User Interface):
HTML5 & CSS3: Used for designing structure and styling, ensuring responsiveness.  
JavaScript (ES6): Handles interactivity, API integration, filters, and dynamic UI updates.

Backend & Database:
Firebase Firestore: NoSQL database used to store editor-contributed articles.  
Firebase Authentication: Manages secure editor login functionality.

API Integration:
NewsAPI: Provides real-time global news data fetched and displayed on the portal.

Cloud Storage (Optional for Extensions):
Firebase Storage: Used for storing images uploaded by editors (optional).

Additional Features & Libraries:
Search & Filters: Implemented in JavaScript for dynamic searching and filtering.  
Responsive Design: CSS media queries ensure compatibility across devices.

• This section must maintain academic tone and be at least 140 words.  
• Use headings exactly like above (Frontend, Backend, API Integration, etc.)  
• Each entry must have a label followed by an explanation.

----------------------------------------------------------
### FEASIBILITY STUDY FORMAT
The feasibility study MUST be divided into the following five sections, written exactly in this style:

1. Technical Feasibility  
Describe the technologies used, their scalability, integration ease, hosting feasibility, etc.

Example structure:
The system uses modern and reliable technologies like Firebase Firestore, Firebase Authentication, and NewsAPI, which are scalable and secure.  
Development with ReactJS/Firebase reduces complexity, etc.

2. Economic Feasibility  
Explain cost benefits, free tiers, affordability, and high cost-to-benefit ratio.

3. Operational Feasibility  
Describe ease of use for editors and readers, low maintenance, smooth workflows, etc.

4. Schedule Feasibility  
Explain timeline feasibility, MVP duration, agile methodology, sprint-based progress, etc.

5. Social Feasibility  
Explain benefits to society, users, digital literacy, and accessibility.

The entire section MUST be at least 250 words and follow the same tone and structure as the example provided.

----------------------------------------------------------
### SYSTEM FEATURES FORMAT
The **system_features** section MUST follow exactly this bullet-style list format:

Editor Authentication: Secure login of editors using Firebase Authentication.  
Editor Dashboard: Provides an interface where editors can add news articles with title, description, and image URL.  
Recently Added Section: Displays editor-contributed news stored in Firebase Firestore.  
API Integration: Fetches real-time news from NewsAPI across categories like Technology, Sports, Entertainment, and Business.  
Search Functionality: Allows users to search specific news articles using keywords.  
Date Filter: Enables filtering by Today, Last 7 Days, or Last 30 Days.  
Responsive Design: Ensures the portal works across desktop, tablet, and mobile devices.  
Interactive UI: Displays news with images, headlines, and descriptions in card format.  
News Details Page: Clicking a card redirects users to a detailed view with full content and related news.  
Secure Content Management: Firebase Firestore ensures secure storage and retrieval of editor-submitted articles.

• Must maintain the same structure and tone.  
• Must be at least 150 words.

----------------------------------------------------------
### MODULES FORMAT
The **modules** section MUST follow this exact structure:

Authentication Module -  
This module manages secure login using Firebase Authentication and ensures that only authorized editors can access the dashboard.

Editor Module -  
Provides an interface for editors to add articles by entering a title, description, and image URL. Validates inputs and stores data in Firestore.

Recently Added Module -  
Displays articles contributed by editors in the “Recently Added” section by fetching them from Firestore in real time.

API News Module -  
Integrates with NewsAPI to fetch real-time news articles across categories such as Politics, Sports, Education, and Technology.

Search Module -  
Allows keyword-based search across both API and editor-contributed articles to enhance accessibility.

Filter Module -  
Enables date-based filtering such as Today, Last 7 Days, and Last 30 Days for efficient browsing.

News Details Module -  
Opens a detailed view of the selected article and fetches related stories for better engagement.

Responsive UI Module -  
Ensures proper layout and component adjustments on desktops, tablets, and mobile devices.

Database Module -  
Manages Firestore operations including adding, reading, and updating articles securely and efficiently.

Scalability Module (Future Scope) -  
A placeholder for future additions such as login for readers, comments, push notifications, AI recommendations, etc.

Requirements:
• Must include all modules listed above.  
• Each module must contain a short explanation.  
• Total section must be at least 300 words.  
• Maintain formal academic tone.


----------------------------------------------------------
### USE CASE FORMAT
The **usecase** section MUST strictly follow the structure below:

Title: <Meaningful title of the use case>

Actors:
User (reader)  
Editor (content contributor)  
Admin (system manager - optional)

Preconditions:
The user has an internet connection.  
The editor has access to the editor dashboard.

Main Flow:
The user opens the portal and views categories of news fetched from the NewsAPI.  
The user browses news cards and clicks on an article to view full details.  
The editor logs into the dashboard and submits a new article with a title, description, and image.  
The submitted article is stored in Firebase Firestore and displayed in the “Recently Added” section.  
The user applies date filters or uses the search bar to find relevant news.  
The system dynamically updates the content.

Postconditions:
Users successfully browse and read news articles.  
Editor-contributed articles appear in the “Recently Added” section in real time.

Requirements:
• Must follow the same headings: Title, Actors, Preconditions, Main Flow, Postconditions.  
• Must be written in clear academic format.  
• Must be at least 120 words.

----------------------------------------------------------
### HARDWARE REQUIREMENTS FORMAT
The hardware requirements section MUST strictly follow the structure and tone shown below.  
It must include hardware requirements for BOTH developers and end users in a clear academic format.  
Total word count (combined with software requirements) must be at least 250 words.

----------------------------------------------------------
Hardware Requirements

1. For Development  
Processor: Intel Core i5 or higher (or equivalent AMD processor)  
RAM: Minimum 8 GB (16 GB recommended for efficient development and multitasking)  
Storage: At least 10 GB of free space for project files, dependencies, SDKs, and build tools  
Operating System:  
• Windows 10 or higher  
• macOS 10.14 (Mojave) or higher  
Screen Resolution: 1366 × 768 or higher (1920 × 1080 recommended for better workspace visibility)  
Internet Connection: Stable broadband connection for NewsAPI usage, package installation, and Firebase integration  

2. For Users (End Users of the Website)  
Device: Desktop, Laptop, Tablet, or Smartphone capable of running a modern web browser  
Processor: Minimum dual-core processor (quad-core recommended for smooth performance)  
RAM: At least 4 GB for comfortable browsing  
Operating System:  
• Windows 8/10/11  
• macOS 10.14 or higher  
• Linux distributions  
• Android 7.0+ or iOS 12+ (for mobile access)  
Internet Connection: Required for retrieving real-time API-based news and loading editor-submitted content from Firestore  

Requirements:  
• Must follow this exact hierarchical structure (1. For Development → hardware list → OS list → resolution → connection → 2. For Users → hardware list → OS list → connection).  
• Must be written in formal academic tone.  
• Must integrate smoothly with the Software Requirements section that follows.

### SOFTWARE REQUIREMENTS FORMAT
The software requirements section MUST strictly follow the structure and tone shown below.  
It must include software requirements for BOTH developers and end users in a clear academic format.  
Total must be at least **250 words** when combined with hardware requirements.

----------------------------------------------------------
Software Requirements

1. For Development  
Languages & Technologies:  
• HTML5, CSS3, JavaScript (ES6)

Frameworks & Libraries:  
• Bootstrap or Custom CSS for responsive design (optional)

Backend & Database:  
• Firebase Firestore (NoSQL cloud database)  
• Firebase Storage (optional – for storing uploaded images)

APIs:  
• NewsAPI (used to fetch real-time global news)

Development Tools:  
• Visual Studio Code (recommended IDE)  

Version Control:  
• Git / GitHub for collaboration and version tracking  

Testing Tools:  
• Chrome DevTools / Firefox Developer Tools for debugging, performance checks, and responsive testing  

2. For End Users  
Operating System:  
• Windows, macOS, Linux, Android, or iOS  

Web Browser:  
• Google Chrome (latest recommended version)  
• Mozilla Firefox (latest version)  
• Microsoft Edge  

Internet Requirement:  
• Active and stable internet connection for fetching API-based news and loading Firestore content  

Requirements:  
• Must follow this exact hierarchical structure (1. For Development → categories → bullet points → descriptions).  
• Must maintain a formal academic tone.  
• Combined with hardware requirements, total must be at least 250 words.

 
 
### OTHER RULES
- Use newline characters (\\n) where appropriate.
- Output must be ONLY the JSON object with NO extra explanation or commentary.
"""

    user_prompt = f"Project Title: {project_title}"
    full_prompt = f"{system_prompt}\n\n{user_prompt}"

    # 4. Call the API
    try:
        print(f"--- Calling Google API for: {project_title} ---")

        response = model.generate_content(full_prompt)

        # Gemini 1.5 returns text directly in response.text
        content_json = json.loads(response.text)

        print("--- Google API Success! ---")
        return content_json

    except Exception as e:
        error_msg = f"ERROR: Google AI content generation failed: {e}"
        print(error_msg)
        if "API key" in str(e):
            error_msg += "\n\nIs your GOOGLE_API_KEY correct and enabled in your Google Cloud project?"
        return {"error": error_msg}


# --- Document Editor Classes ---

class ProjectDocumentEditor:
    def __init__(self, doc_path):
        self.doc_path = doc_path
        self.replacements = 0
        self.changes_log = []

    def process_document(self, old_data_dict, new_data_dict, uploaded_images=None, screenshot_metadata=None):
        with open(self.doc_path, 'rb') as docx_file:
            self.doc = Document(docx_file)

        for field_type, old_value in old_data_dict.items():
            if old_value and new_data_dict.get(field_type):
                # --- START: FIX ---
                # Skip the dedicated screenshot tag in the generic text replacement loop.
                # This ensures it's only handled by the _replace_screenshots method.
                if old_value == SCREENSHOT_BLOCK_TAG:
                    continue
                # --- END: FIX ---
                new_value = new_data_dict[field_type]
                self._replace_text_globally(old_value, new_value, field_type)

        if uploaded_images:
            self._replace_images(uploaded_images)

        # Process Screenshots using its dedicated function
        if screenshot_metadata and uploaded_images:
            self._replace_screenshots(uploaded_images, screenshot_metadata)

        return self.doc, self.replacements, self.changes_log

    def replace_project_details(self, old_data_dict, new_data_dict, uploaded_images=None, screenshot_metadata=None, output_path=None):
        if output_path is None:
            base_name = os.path.splitext(self.doc_path)[0]
            output_path = f"{base_name}_COMPLETE_UPDATE.docx"

        doc, replacements, changes_log = self.process_document(old_data_dict, new_data_dict, uploaded_images, screenshot_metadata)

        self._add_field_update_trigger()

        doc.save(output_path)
        return replacements, changes_log, output_path

    def get_document_as_text(self, old_data_dict, new_data_dict):
        doc, _, _ = self.process_document(old_data_dict, new_data_dict)
        return "\n".join([p.text for p in doc.paragraphs])

    def _add_field_update_trigger(self):
        """Adds XML to the document settings to force field (TOC) update upon opening."""
        from lxml import etree
        settings_part = self.doc.settings.element
        update_fields_tag = qn('w:updateFields')

        if settings_part.find(update_fields_tag) is None:
            new_field_update_element = etree.Element(update_fields_tag)
            new_field_update_element.set(qn('w:val'), 'true')
            settings_part.append(new_field_update_element)
            self.changes_log.append("SYSTEM: Field update (TOC refresh) tag added to document settings.")

    def _replace_code_list(self, code_json_string):
        """
        Parses JSON of code files and inserts them sequentially at the placeholder's location.
        Implements required formatting: uppercase filename, no 'File' prefix, 14pt size.
        """

        try:
            code_blocks = json.loads(code_json_string)
        except json.JSONDecodeError:
            self.changes_log.append("ERROR: Failed to parse code blocks JSON.")
            return

        if not code_blocks:
            return

        for paragraph in self.doc.paragraphs:
            if CODE_LIST_TAG in paragraph.text:

                parent_element = paragraph._element.getparent()
                anchor_index = parent_element.index(paragraph._element)

                # We iterate backward to stack items correctly at the fixed anchor_index
                for code_data in reversed(code_blocks):
                    filename = code_data['filename'].upper()
                    content = code_data['content']

                    # 1. Prepare Code Content Paragraph (p_code)
                    p_code = self.doc.add_paragraph()
                    p_code.style = 'Normal'

                    for j, line in enumerate(content.split('\n')):
                        run = p_code.add_run(line)
                        try:
                            run.style = 'Code' # Attempt to apply a monospace/code style
                        except:
                            pass

                        if j < len(content.split('\n')) - 1:
                            run.add_break()

                    # 2. Prepare Filename Paragraph (p_filename)
                    p_filename = self.doc.add_paragraph()
                    filename_run = p_filename.add_run(filename)

                    # Set font size to 14 points and make bold (block format)
                    filename_run.font.size = Pt(14)
                    filename_run.bold = True

                    # --- Targeted XML Insertion (Insert both at the same index) ---

                    parent_element.insert(anchor_index, p_code._element)
                    parent_element.insert(anchor_index, p_filename._element)

                    self.changes_log.append(f"CODE_BLOCK added: {filename}")
                    self.replacements += 2

                # 3. Delete the original placeholder paragraph last
                parent_element.remove(paragraph._element)

                return

    def _replace_screenshots(self, uploaded_images, screenshot_metadata):
        """Inserts named screenshots into the document."""

        # screenshot_metadata is a list of {name: "Screen Name", file_id: "screenshot_file_1"}
        if not screenshot_metadata:
            return

        for paragraph in self.doc.paragraphs:
            if SCREENSHOT_BLOCK_TAG in paragraph.text:

                parent_element = paragraph._element.getparent()
                anchor_index = parent_element.index(paragraph._element)

                # Iterate backward to stack correctly
                for ss_data in reversed(screenshot_metadata):
                    ss_file_id = ss_data['file_id']
                    ss_name = ss_data['name'].upper()

                    # Determine the actual saved filename
                    saved_filename = uploaded_images.get(ss_file_id)

                    if not saved_filename:
                        self.changes_log.append(f"WARNING: Screenshot file missing for {ss_name}.")
                        continue

                    image_path = os.path.join(app.config['UPLOAD_FOLDER'], saved_filename)

                    # 1. Insert Name Header
                    p_name = self.doc.add_paragraph()
                    p_name.add_run(ss_name).bold = True
                    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # 2. Insert Image
                    p_image = self.doc.add_paragraph()
                    run_image = p_image.add_run()
                    run_image.add_picture(image_path, width=Inches(6.0))
                    p_image.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_image.paragraph_format.space_after = Inches(0.2) # Add spacing after image

                    # Targeted XML Insertion (Insert both at the same index)
                    parent_element.insert(anchor_index, p_image._element)
                    parent_element.insert(anchor_index, p_name._element)

                    self.changes_log.append(f"SCREENSHOT added: {ss_name}")
                    self.replacements += 2

                # Delete the original placeholder paragraph last
                parent_element.remove(paragraph._element)
                return

    def _replace_text_globally(self, old_text, new_text, field_type):
        # 1. Check for the single CODE LIST TAG
        if old_text == CODE_LIST_TAG:
            self._replace_code_list(new_text)
            return

        # 2. Standard Text Replacement (Original Logic)
        for i, paragraph in enumerate(self.doc.paragraphs):
            if paragraph.text.strip():
                self._smart_replace_with_formatting(paragraph, old_text, new_text, f"paragraph_{i+1}", field_type)

        for table_idx, table in enumerate(self.doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        if paragraph.text.strip():
                            location = f"table_{table_idx+1}.cell_{cell_idx+1}.para_{para_idx+1}"
                            self._smart_replace_with_formatting(paragraph, old_text, new_text, location, field_type)

        for section_idx, section in enumerate(self.doc.sections):
            if section.header is not None:
                for para_idx, paragraph in enumerate(section.header.paragraphs):
                    if paragraph.text.strip():
                        location = f"header_section_{section_idx+1}.para_{para_idx+1}"
                        self._smart_replace_with_formatting(paragraph, old_text, new_text, location, field_type)

            if section.footer is not None:
                for para_idx, paragraph in enumerate(section.footer.paragraphs):
                    if paragraph.text.strip():
                        location = f"footer_section_{section_idx+1}.para_{para_idx+1}"
                        self._smart_replace_with_formatting(paragraph, old_text, new_text, location, field_type)

    def _replace_images(self, uploaded_images):

        paragraphs_to_check = list(self.doc.paragraphs)

        for paragraph in paragraphs_to_check:
            for placeholder, filename in uploaded_images.items():
                tag = f"[[{placeholder}]]"

                if tag in paragraph.text:
                    try:
                        image_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)

                        section = self.doc.sections[0]
                        page_height_emu = section.page_height.emu
                        margin_top_emu = section.top_margin.emu
                        margin_bottom_emu = section.bottom_margin.emu
                        usable_height_emu_value = page_height_emu - margin_top_emu - margin_bottom_emu
                        usable_height_inches = Length(usable_height_emu_value).inches
                        MAX_HEIGHT_INCHES = usable_height_inches * 0.85

                        img = Image.open(image_path)
                        img_width, img_height = img.size
                        aspect_ratio = img_width / img_height
                        MAX_WIDTH_INCHES = 6.0

                        final_height_inches = MAX_HEIGHT_INCHES
                        final_width_inches = MAX_HEIGHT_INCHES * aspect_ratio

                        if final_width_inches > MAX_WIDTH_INCHES:
                             final_width_inches = MAX_WIDTH_INCHES
                             final_height_inches = MAX_WIDTH_INCHES / aspect_ratio

                        parent = paragraph._element.getparent()
                        new_p_element = paragraph._element.makeelement(qn('w:p'))
                        parent.insert(parent.index(paragraph._element) + 1, new_p_element)

                        new_paragraph = Paragraph(new_p_element, self.doc)

                        new_paragraph.style = 'Normal'
                        new_paragraph.paragraph_format.space_before = 0
                        new_paragraph.paragraph_format.space_after = 0

                        run = new_paragraph.add_run()
                        run.add_picture(image_path, width=Inches(final_width_inches), height=Inches(final_height_inches))
                        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        paragraph._element.getparent().remove(paragraph._element)

                        self.replacements += 1
                        log_entry = f"IMAGE in {placeholder}: '{tag}' replaced with '{filename}' (Sized: {final_width_inches:.2f}x{final_height_inches:.2f} in)"
                        self.changes_log.append(log_entry)

                    except Exception as e:
                        error_message = f"ERROR: Image replacement failed for {placeholder}. Details: {e}"
                        self.changes_log.append(error_message)
                        print(error_message)

    def _smart_replace_with_formatting(self, paragraph, old_text, new_text, location, field_type):
        original_text = paragraph.text
        if not original_text or old_text.lower() not in original_text.lower():
            return

        original_runs = []
        for run in paragraph.runs:
            original_runs.append({
                'text': run.text, 'font_name': run.font.name, 'font_size': run.font.size,
                'bold': run.bold, 'italic': run.italic, 'underline': run.underline,
                'color': run.font.color.rgb if run.font.color else None, 'all_caps': run.font.all_caps
            })

        paragraph.clear()

        pattern = re.compile(f'({re.escape(old_text)})', re.IGNORECASE)
        parts = pattern.split(original_text)

        current_pos = 0
        for part in parts:
            if not part: continue
            if part.lower() == old_text.lower():
                formatting = self._get_formatting_at_position(original_runs, current_pos)
                run = paragraph.add_run(new_text)
                self._apply_formatting(run, formatting)
                self.replacements += 1
                log_entry = f"{field_type.upper()} in {location}: '{part}' → '{new_text}'"
                self.changes_log.append(log_entry)
            else:
                self._add_text_with_original_formatting(paragraph, part, original_runs, current_pos)
            current_pos += len(part)

    def _get_formatting_at_position(self, runs_data, position):
        current_pos = 0
        for run_data in runs_data:
            run_length = len(run_data['text'])
            if position < current_pos + run_length:
                return {
                    'font_name': run_data['font_name'], 'font_size': run_data['font_size'],
                    'bold': run_data['bold'], 'italic': run_data['italic'], 'underline': run_data['underline'],
                    'color': run_data['color'], 'all_caps': run_data['all_caps']
                }
            current_pos += run_length
        return None

    def _add_text_with_original_formatting(self, paragraph, text, runs_data, start_pos):
        for i, char in enumerate(text):
            pos = start_pos + i
            formatting = self._get_formatting_at_position(runs_data, pos)
            run = paragraph.add_run(char)
            self._apply_formatting(run, formatting)

    def _apply_formatting(self, run, formatting):
        if formatting:
            if formatting['font_name']:
                run.font.name = formatting['font_name']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), formatting['font_name'])
            if formatting['font_size']:
                run.font.size = formatting['font_size']
            if formatting['bold'] is not None:
                run.bold = formatting['bold']
            if formatting['italic'] is not None:
                run.italic = formatting['italic']
            if formatting['underline'] is not None:
                run.underline = formatting['underline']
            if formatting['color']:
                run.font.color.rgb = formatting['color']
            if formatting['all_caps'] is not None:
                run.font.all_caps = formatting['all_caps']

class SmartProjectBookEditor:
    def __init__(self, doc_path):
        self.doc_path = doc_path

    def complete_update(self, updates_dict, uploaded_images, screenshot_metadata=None):
        editor = ProjectDocumentEditor(self.doc_path)

        old_data = {}
        new_data = {}

        # Include 'screenshot_block' in the list of potential fields
        fields_to_process = ['name', 'project', 'professor', 'guide', 'year'] + DOCUMENTATION_TAGS + ['code_list', 'screenshot_block']

        for field in fields_to_process:
            if f'old_{field}' in updates_dict and f'new_{field}' in updates_dict:
                old_data[field] = updates_dict[f'old_{field}']
                new_data[field] = updates_dict[f'new_{field}']

        return editor.replace_project_details(old_data, new_data, uploaded_images, screenshot_metadata)

    def get_preview_text(self, updates_dict):
        editor = ProjectDocumentEditor(self.doc_path)

        old_data = {}
        new_data = {}

        for field in ['name', 'project', 'professor', 'guide', 'year'] + DOCUMENTATION_TAGS + ['code_list']:
            if f'old_{field}' in updates_dict and f'new_{field}' in updates_dict:
                old_data[field] = updates_dict[f'old_{field}']
                new_data[field] = updates_dict[f'new_{field}']

        return editor.get_document_as_text(old_data, new_data)


# --- Flask Routes (The rest of the file remains the same) ---
@app.route('/', methods=['GET'])
def index():
    if not TEMPLATE_CONFIGS:
        return "Error: No templates found in templates.json.", 500

    default_config = TEMPLATE_CONFIGS[0]

    current_year = datetime.datetime.now().year
    next_year = current_year + 1
    default_new_year = f"{current_year}-{next_year}"

    return render_template(
        'index.html',
        template_configs=TEMPLATE_CONFIGS,
        old_details=default_config['old_details'],
        default_template_file=default_config['file'],
        default_new_year=default_new_year
    )

@app.route('/preview', methods=['POST'])
def preview():
    template_file = request.json.get('template_file')
    selected_config = next((c for c in TEMPLATE_CONFIGS if c['file'] == template_file), None)

    if not selected_config:
        return jsonify({"error": "Invalid template selected."}), 400

    old_details = selected_config['old_details']

    new_student_name = request.json.get('student_name')
    new_project_title = request.json.get('project_title')
    new_professor_name = request.json.get('professor_name')
    new_guide_name = request.json.get('guide_name')
    new_year = request.json.get('year')
    new_code_list = request.json.get('all_code_data')

    updates = {
        'old_name': old_details['student_name'],
        'new_name': new_student_name or old_details['student_name'],
        'old_project': old_details['project_title'],
        'new_project': new_project_title or old_details['project_title'],
        'old_professor': old_details['professor_name'],
        'new_professor': new_professor_name or old_details['professor_name'],
        'old_guide': old_details['guide_name'],
        'new_guide': new_guide_name or old_details['guide_name'],
        'old_year': old_details['year'],
        'new_year': new_year or old_details['year'],
        'new_code_list': new_code_list # Pass code list for preview
    }

    # Add documentation content from form
    for tag in DOCUMENTATION_TAGS:
        new_content = request.json.get(tag)
        if new_content:
            updates[f'old_{tag}'] = f"[[{tag}]]"
            updates[f'new_{tag}'] = new_content

    doc_path = os.path.join(os.getcwd(), template_file)
    if not os.path.exists(doc_path):
        return jsonify({"error": f"Template file '{template_file}' not found."}), 404

    editor = SmartProjectBookEditor(doc_path)
    updated_text = editor.get_preview_text(updates)

    preview_html = updated_text.replace('\n', '<br>')

    return jsonify({"html": preview_html})

@app.route('/upload_and_update', methods=['POST'])
def upload_and_update():
    # 1. Retrieve Template Info and set up file handling
    template_file = request.form.get('template_file')
    selected_config = next((c for c in TEMPLATE_CONFIGS if c['file'] == template_file), None)

    if not selected_config:
        return "Error: Invalid template selected.", 400

    old_details = selected_config['old_details']
    doc_path = os.path.join(os.getcwd(), template_file)

    # 2. Handle File Uploads (Manual or Generated)
    uploaded_files = {}
    image_types = [
        'plagarism_report', 'er_diagram', 'class_diagram', 'activity_diagram',
        'usecase_diagram', 'sequence_diagram', 'component_diagram', 'deployment_diagram', 'database_model'
    ]

    generated_data_json = request.form.get('generated_files_data')
    if generated_data_json:
        try:
            generated_files = json.loads(generated_data_json)
            uploaded_files.update(generated_files)
        except json.JSONDecodeError:
            print("ERROR: Could not decode generated files JSON.")
            pass

    for img_type in image_types:
        if img_type in request.files:
            file = request.files[img_type]
            if file and file.filename != '':
                filename = f"{int(time.time())}_{file.filename}"
                file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                uploaded_files[img_type] = filename

    # 2a. Handle Screenshot Uploads
    screenshot_metadata = []
    screenshot_data_json = request.form.get('all_screenshot_data')
    if screenshot_data_json:
        try:
            screenshot_metadata = json.loads(screenshot_data_json)
            # Save the actual files and map their file_id to the saved filename
            for ss_item in screenshot_metadata:
                file_id = ss_item.get('file_id') # e.g., 'screenshot_file_1'
                if file_id in request.files:
                    file = request.files[file_id]
                    if file and file.filename != '':
                        # Create a unique filename and save the file
                        filename = f"{int(time.time())}_{file.filename}"
                        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                        # Add the file_id -> saved_filename mapping for the doc editor
                        uploaded_files[file_id] = filename
        except json.JSONDecodeError:
            print("ERROR: Could not decode screenshot metadata JSON.")
            pass # Fail silently or return an error


    # 3. Extract Text Data and Prepare Updates
    new_student_name = request.form.get('student_name')
    new_project_title = request.form.get('project_title')
    new_professor_name = request.form.get('professor_name')
    new_guide_name = request.form.get('guide_name')
    new_year = request.form.get('year')
    new_code_list = request.form.get('all_code_data') # Get JSON code list

    # Base updates dictionary
    updates = {
        'old_name': old_details['student_name'],
        'new_name': new_student_name,
        'old_project': old_details['project_title'],
        'new_project': new_project_title,
        'old_professor': old_details['professor_name'],
        'new_professor': new_professor_name,
        'old_guide': old_details['guide_name'],
        'new_guide': new_guide_name,
        'old_year': old_details['year'],
        'new_year': new_year
    }

    # Add documentation content from form
    for tag in DOCUMENTATION_TAGS:
        new_content = request.form.get(tag)
        if new_content:
            updates[f'old_{tag}'] = f"[[{tag}]]"
            updates[f'new_{tag}'] = new_content

    # Add the CODE LIST tag replacement pair
    if new_code_list and json.loads(new_code_list):
        updates['old_code_list'] = CODE_LIST_TAG
        updates['new_code_list'] = new_code_list

    # Add the SCREENSHOT BLOCK tag replacement pair
    if screenshot_metadata:
        # This old/new pair is used to find the right keys in the `complete_update` function.
        updates['old_screenshot_block'] = SCREENSHOT_BLOCK_TAG
        updates['new_screenshot_block'] = "[Processed dynamically]"


    # 4. Filter for Final Updates
    final_updates = {}

    for key, value in updates.items():
        if key.startswith('old_')  and value:
            new_key = 'new_' + key[4:]

            # Check if the field is a special content block
            is_content_block = any(key.endswith(f"_{t}") for t in DOCUMENTATION_TAGS) or key.endswith('_code_list') or key.endswith('_screenshot_block')

            if updates.get(new_key) and (updates.get(new_key) != value or is_content_block):
                final_updates[key] = value
                final_updates[new_key] = updates[new_key]

    # 5. Process the Document
    if not os.path.exists(doc_path):
        return f"Error: Template file '{template_file}' not found.", 404

    editor = SmartProjectBookEditor(doc_path)
    # Pass the populated screenshot_metadata to the editor
    replacements, changes_log, output_path = editor.complete_update(
        final_updates,
        uploaded_files,
        screenshot_metadata=screenshot_metadata
    )

    return jsonify({
        "status": "success",
        "replacements": replacements,
        "changes_log": changes_log,
        "output_path": output_path
    })
@app.route('/generate_project_content', methods=['POST'])
def generate_project_content():
    """
    This route is called by the new 'Generate with AI' button.
    It gets the project title, calls the AI helper, and returns
    the generated content as JSON.
    """
    project_title = request.json.get('project_title')
    
    if not project_title:
        return jsonify({"error": "Project title is required."}), 400

    # Call our new helper function
    content_data = get_ai_generated_documentation(project_title)

    if "error" in content_data:
         return jsonify(content_data), 500

    # Success! Send the JSON data back to the webpage
    return jsonify(content_data)

@app.route('/generate_diagrams', methods=['POST'])
def generate_diagrams():
    pass

if __name__=="__main__":
    app.run(debug=True)