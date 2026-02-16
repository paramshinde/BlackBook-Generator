import os
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


def build_docx(document_json: dict, template_profile: dict, output_path: str) -> str:
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    margins = template_profile.get("margins", {})
    if margins:
        # Accept inches as numeric values for API simplicity.
        section.top_margin = Inches(float(margins.get("top", 1.0)))
        section.bottom_margin = Inches(float(margins.get("bottom", 1.0)))
        section.left_margin = Inches(float(margins.get("left", 1.0)))
        section.right_margin = Inches(float(margins.get("right", 1.0)))

    title = document_json.get("meta", {}).get("projectTitle", "Black Book")
    doc.add_heading(title, level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)
    for toc in document_json.get("toc", []):
        doc.add_paragraph(f"{toc.get('title', '')} .... {toc.get('page', '')}")

    for section_data in document_json.get("sections", []):
        doc.add_heading(section_data.get("heading", "Section"), level=1)
        doc.add_paragraph(section_data.get("content", ""))

    for i, figure in enumerate(document_json.get("figures", []), start=1):
        doc.add_paragraph(f"Figure {i}: {figure.get('caption', 'Screenshot')}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    return output_path


def build_docx_bytes(document_json: dict, template_profile: dict) -> bytes:
    temp = BytesIO()
    doc = Document()
    doc.add_heading(document_json.get("meta", {}).get("projectTitle", "Black Book"), 0)
    doc.save(temp)
    return temp.getvalue()
