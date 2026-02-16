import os
import uuid
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image

from config import GENERATED_DIR, MAX_IMAGE_HEIGHT_INCHES, MAX_IMAGE_WIDTH_INCHES
from services.preview_service import DOC_KEYS

CODE_PLACEHOLDER = "[[CODE_IMPLEMENTATION_BLOCK]]"
SCREENSHOT_PLACEHOLDER = "[[SCREENSHOT_BLOCK]]"
DIAGRAM_PLACEHOLDER = "[[DIAGRAM_BLOCK]]"
PLAGIARISM_PLACEHOLDER = "[[plagiarism_report]]"
SCREENSHOT_PLACEHOLDERS = {
    "[[SCREENSHOT_BLOCK]]",
    "[[screenshot_block]]",
}
DIAGRAM_PLACEHOLDERS = {
    "[[DIAGRAM_BLOCK]]",
    "[[diagram_block]]",
    "[[event_table]]",
    "[[er_diagram]]",
    "[[class_diagram]]",
    "[[activity_diagram]]",
    "[[use_case_diagram]]",
    "[[usecase_diagram]]",
    "[[sequence_diagram]]",
    "[[component_diagram]]",
    "[[deployment_diagram]]",
    "[[database_model]]",
}
PLAGIARISM_PLACEHOLDERS = {
    "[[plagiarism_report]]",
    "[[plagarism_report]]",
    "[[PLAGIARISM_REPORT]]",
    "[[PLAGARISM_REPORT]]",
}


def _sanitize_xml_text(value):
    if value is None:
        return ""
    text = str(value)
    # XML 1.0 valid chars: tab, LF, CR, and U+0020..U+D7FF, U+E000..U+FFFD.
    return "".join(
        ch for ch in text if ch in ("\t", "\n", "\r") or (0x20 <= ord(ch) <= 0xD7FF) or (0xE000 <= ord(ch) <= 0xFFFD)
    )


def _replace_text_paragraph(paragraph, old_text: str, new_text: str):
    if not old_text:
        return False
    new_text = _sanitize_xml_text(new_text)

    # Prefer in-run replacement to preserve template formatting.
    replaced = False
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            replaced = True

    if replaced:
        return True

    # Handle placeholders split across multiple runs while keeping the first run style.
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text in full_text and paragraph.runs:
        updated = full_text.replace(old_text, new_text)
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
        return True

    if old_text in paragraph.text:
        paragraph.text = paragraph.text.replace(old_text, new_text)
        return True

    return False


def _replace_text_globally(doc: Document, old_text: str, new_text: str):
    for paragraph in doc.paragraphs:
        _replace_text_paragraph(paragraph, old_text, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_text_paragraph(paragraph, old_text, new_text)


def _set_run_font(run, size=12, bold=False):
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)


def _apply_heading_style(paragraph, level=1):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        paragraph.style = "Heading 1"
        for run in paragraph.runs:
            _set_run_font(run, size=16, bold=True)
    else:
        paragraph.style = "Heading 2"
        for run in paragraph.runs:
            _set_run_font(run, size=14, bold=True)


def _apply_body_style(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(8)
    for run in paragraph.runs:
        _set_run_font(run, size=12)


def _resize_image(path: str):
    img = Image.open(path)
    w, h = img.size
    if w == 0 or h == 0:
        return MAX_IMAGE_WIDTH_INCHES, MAX_IMAGE_HEIGHT_INCHES

    aspect = w / h
    width = MAX_IMAGE_WIDTH_INCHES
    height = width / aspect

    if height > MAX_IMAGE_HEIGHT_INCHES:
        height = MAX_IMAGE_HEIGHT_INCHES
        width = height * aspect

    return width, height


def _add_figure(paragraph, image_path, caption_text, figure_no):
    parent = paragraph._element.getparent()
    anchor_index = parent.index(paragraph._element)

    p_img = paragraph._parent.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    width, height = _resize_image(image_path)
    run_img = p_img.add_run()
    run_img.add_picture(image_path, width=Inches(width), height=Inches(height))

    p_caption = paragraph._parent.add_paragraph(f"Figure {figure_no}: {_sanitize_xml_text(caption_text)}")
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_caption.runs:
        _set_run_font(run, size=11, bold=False)

    parent.insert(anchor_index + 1, p_img._element)
    parent.insert(anchor_index + 2, p_caption._element)


def _add_screenshot_block(paragraph, image_path, title_text):
    parent = paragraph._element.getparent()
    anchor_index = parent.index(paragraph._element)

    p_title = paragraph._parent.add_paragraph(_sanitize_xml_text(title_text or "Screenshot"))
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p_title.runs:
        _set_run_font(run, size=12, bold=True)

    p_img = paragraph._parent.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    width, height = _resize_image(image_path)
    run_img = p_img.add_run()
    run_img.add_picture(image_path, width=Inches(width), height=Inches(height))

    parent.insert(anchor_index + 1, p_title._element)
    parent.insert(anchor_index + 2, p_img._element)


def _add_insert_error_note(paragraph, message):
    parent = paragraph._element.getparent()
    anchor_index = parent.index(paragraph._element)
    p_note = paragraph._parent.add_paragraph(_sanitize_xml_text(message))
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_note.runs:
        _set_run_font(run, size=10, bold=False)
    parent.insert(anchor_index + 1, p_note._element)


def _add_code_block_after(paragraph, code_items):
    parent = paragraph._element.getparent()
    anchor_index = parent.index(paragraph._element)

    for item in code_items:
        p_name = paragraph._parent.add_paragraph(_sanitize_xml_text(item["filename"]).upper())
        p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p_name.runs:
            _set_run_font(run, size=12, bold=True)

        p_code = paragraph._parent.add_paragraph(_sanitize_xml_text(item["content"]))
        p_code.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        for run in p_code.runs:
            _set_run_font(run, size=10, bold=False)

        parent.insert(anchor_index + 1, p_name._element)
        parent.insert(anchor_index + 2, p_code._element)
        anchor_index += 2


def _normalize_diagram_text(value: str) -> str:
    return "".join(ch.lower() if ch.isalnum() else " " for ch in (value or "")).strip()


def _match_diagram_index_for_placeholder(diagrams, placeholder_token: str, used_indexes: set[int]) -> Optional[int]:
    token = placeholder_token.replace("[[", "").replace("]]", "").strip().lower()
    aliases = {
        "event_table": ["event table", "event_table", "eventtable"],
        "er_diagram": ["entity relationship", "er diagram", "er_diagram", "erd"],
        "class_diagram": ["class diagram", "class_diagram"],
        "activity_diagram": ["activity diagram", "activity_diagram"],
        "use_case_diagram": ["use case diagram", "use_case_diagram", "usecase diagram", "usecase_diagram"],
        "usecase_diagram": ["use case diagram", "use_case_diagram", "usecase diagram", "usecase_diagram"],
        "sequence_diagram": ["sequence diagram", "sequence_diagram"],
        "component_diagram": ["component diagram", "component_diagram"],
        "deployment_diagram": ["deployment diagram", "deployment_diagram"],
        "database_model": ["database model", "database diagram", "database_model"],
    }
    keywords = aliases.get(token, [token.replace("_", " ")])

    for idx, item in enumerate(diagrams):
        if idx in used_indexes:
            continue
        name = _normalize_diagram_text(item.get("name", ""))
        path_name = _normalize_diagram_text(os.path.basename(item.get("path", "")))
        haystack = f"{name} {path_name}"
        if any(_normalize_diagram_text(k) in haystack for k in keywords):
            return idx
    return None


def _inject_blocks(doc: Document, code_items, screenshots, diagrams, plagiarism_report):
    figure_count = 1
    inserted_screenshots = False
    inserted_diagrams = False
    inserted_plagiarism = False
    used_diagram_indexes = set()

    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        text_lower = text.lower()

        if CODE_PLACEHOLDER in text:
            _add_code_block_after(paragraph, code_items)
            paragraph.text = paragraph.text.replace(CODE_PLACEHOLDER, "")

        screenshot_hit = any(token.lower() in text_lower for token in SCREENSHOT_PLACEHOLDERS)
        if screenshot_hit:
            for item in screenshots:
                try:
                    _add_screenshot_block(paragraph, item["path"], item["name"])
                    inserted_screenshots = True
                except Exception:
                    _add_insert_error_note(paragraph, f"[Image skipped: {item.get('name', 'screenshot')}]")
            for token in SCREENSHOT_PLACEHOLDERS:
                paragraph.text = paragraph.text.replace(token, "")

        diagram_hit = any(token.lower() in text_lower for token in DIAGRAM_PLACEHOLDERS)
        if diagram_hit:
            matched_tokens = [token for token in DIAGRAM_PLACEHOLDERS if token.lower() in text_lower]
            generic_tokens = {"[[diagram_block]]"}
            specific_tokens = [token for token in matched_tokens if token.lower() not in generic_tokens]

            target_indexes = []
            if specific_tokens:
                for token in specific_tokens:
                    match_idx = _match_diagram_index_for_placeholder(diagrams, token, used_diagram_indexes)
                    if match_idx is not None:
                        target_indexes.append(match_idx)
            else:
                target_indexes = [idx for idx in range(len(diagrams)) if idx not in used_diagram_indexes]

            for idx in target_indexes:
                item = diagrams[idx]
                p_title = paragraph._parent.add_paragraph(_sanitize_xml_text(item["name"]))
                p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p_title.runs:
                    _set_run_font(run, size=12, bold=True)
                paragraph._element.getparent().insert(paragraph._element.getparent().index(paragraph._element) + 1, p_title._element)
                try:
                    _add_figure(paragraph, item["path"], item["name"], figure_count)
                    figure_count += 1
                    inserted_diagrams = True
                    used_diagram_indexes.add(idx)
                except Exception:
                    _add_insert_error_note(paragraph, f"[Diagram skipped: {item.get('name', 'diagram')}]")
            for token in DIAGRAM_PLACEHOLDERS:
                paragraph.text = paragraph.text.replace(token, "")

        plagiarism_hit = any(token.lower() in text_lower for token in PLAGIARISM_PLACEHOLDERS)
        if plagiarism_hit:
            if plagiarism_report and plagiarism_report.get("path"):
                try:
                    _add_figure(paragraph, plagiarism_report["path"], plagiarism_report.get("title", "Plagiarism Report"), figure_count)
                    figure_count += 1
                    inserted_plagiarism = True
                except Exception:
                    _add_insert_error_note(paragraph, "[Plagiarism report image skipped]")
            elif plagiarism_report and plagiarism_report.get("name"):
                _add_insert_error_note(paragraph, f"[Plagiarism report attached: {plagiarism_report['name']}]")
            for token in PLAGIARISM_PLACEHOLDERS:
                paragraph.text = paragraph.text.replace(token, "")

    # Fallback: if a template has no matching placeholder, still append uploaded media.
    if screenshots and not inserted_screenshots:
        doc.add_paragraph("Screenshots")
        anchor = doc.paragraphs[-1]
        for item in screenshots:
            try:
                _add_screenshot_block(anchor, item["path"], item["name"])
            except Exception:
                _add_insert_error_note(anchor, f"[Image skipped: {item.get('name', 'screenshot')}]")

    if diagrams and not inserted_diagrams:
        doc.add_paragraph("Diagrams")
        anchor = doc.paragraphs[-1]
        for idx, item in enumerate(diagrams):
            try:
                _add_figure(anchor, item["path"], item["name"], figure_count)
                figure_count += 1
                used_diagram_indexes.add(idx)
            except Exception:
                _add_insert_error_note(anchor, f"[Diagram skipped: {item.get('name', 'diagram')}]")

    if plagiarism_report and not inserted_plagiarism and plagiarism_report.get("path"):
        doc.add_paragraph(plagiarism_report.get("title", "Plagiarism Report"))
        anchor = doc.paragraphs[-1]
        try:
            _add_figure(anchor, plagiarism_report["path"], plagiarism_report.get("title", "Plagiarism Report"), figure_count)
        except Exception:
            _add_insert_error_note(anchor, "[Plagiarism report image skipped]")


def _add_page_numbers(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    p._element.append(fld_begin)
    p._element.append(instr)
    p._element.append(fld_end)


def _insert_toc_field(doc: Document):
    toc_heading = doc.add_paragraph("Table of Contents")
    _apply_heading_style(toc_heading, level=1)

    p = doc.add_paragraph()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "Update field to populate TOC."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r = OxmlElement("w:r")
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_text)
    r.append(fld_end)
    p._element.append(r)


def _set_document_layout(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def _style_document(doc: Document):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            if paragraph.style.name == "Heading 1":
                _apply_heading_style(paragraph, level=1)
            else:
                _apply_heading_style(paragraph, level=2)
        elif len(text) < 80 and text.isupper():
            _apply_heading_style(paragraph, level=2)
        else:
            _apply_body_style(paragraph)


def _resolve_plagiarism_path(plagiarism_report):
    if not plagiarism_report or not plagiarism_report.get("path"):
        return plagiarism_report

    ext = os.path.splitext(plagiarism_report["path"])[1].lower()
    if ext == ".pdf":
        # Keep generation resilient by skipping in-process PDF rendering.
        plagiarism_report["path"] = None

    return plagiarism_report


def _append_references_section(doc: Document, references):
    if not isinstance(references, list):
        return

    normalized = []
    for item in references:
        if isinstance(item, dict):
            label = _sanitize_xml_text(item.get("label", "")).strip()
            url = _sanitize_xml_text(item.get("url", "")).strip()
        else:
            label = ""
            url = _sanitize_xml_text(item).strip()
        if url:
            normalized.append({"label": label, "url": url})

    if not normalized:
        return

    doc.add_paragraph("References")
    for paragraph in doc.paragraphs[-1:]:
        for run in paragraph.runs:
            _set_run_font(run, size=14, bold=True)

    for idx, ref in enumerate(normalized, start=1):
        text = f"{idx}. {ref['label']} - {ref['url']}" if ref["label"] else f"{idx}. {ref['url']}"
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            _set_run_font(run, size=11, bold=False)


def generate_document(
    template_path: str,
    template_config: dict,
    payload: dict,
    code_items: list,
    screenshots: list,
    diagrams: list,
    plagiarism_report: dict | None = None,
):
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)

    student = payload.get("studentDetails", {})
    documentation = payload.get("documentation", {})
    old = template_config.get("old_details", {})

    replacements = [
        (old.get("student_name", "[[student_name]]"), _sanitize_xml_text(student.get("name", ""))),
        (old.get("project_title", "[[project_title]]"), _sanitize_xml_text(student.get("project", ""))),
        (old.get("professor_name", "[[professor_name]]"), _sanitize_xml_text(student.get("professor", ""))),
        (old.get("guide_name", "[[guide_name]]"), _sanitize_xml_text(student.get("guide", ""))),
        (old.get("year", "[[year]]"), _sanitize_xml_text(student.get("year", ""))),
    ]

    for key in DOC_KEYS:
        old_value = old.get(key, f"[[{key}]]")
        replacements.append((old_value, _sanitize_xml_text(documentation.get(key, ""))))

    for old_text, new_text in replacements:
        if old_text and new_text is not None:
            _replace_text_globally(doc, str(old_text), str(new_text))

    plagiarism_report = _resolve_plagiarism_path(plagiarism_report)
    _inject_blocks(doc, code_items, screenshots, diagrams, plagiarism_report)

    _insert_toc_field(doc)
    _add_page_numbers(doc)
    _append_references_section(doc, payload.get("references", []))

    Path(GENERATED_DIR).mkdir(parents=True, exist_ok=True)
    out_name = f"blackbook_{uuid.uuid4().hex[:10]}.docx"
    out_path = os.path.join(GENERATED_DIR, out_name)
    doc.save(out_path)
    return out_name, out_path
