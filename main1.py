from docx import Document
import re
import os
from docx.oxml.ns import qn

class ProjectDocumentEditor:
    def __init__(self, doc_path):
        self.doc_path = doc_path
        self.doc = Document(doc_path)
        self.replacements = 0
        self.changes_log = []
    
    def replace_project_details(self, old_data_dict, new_data_dict, output_path=None):
        """
        Replace multiple project details throughout the document
        
        Args:
            old_data_dict: {'name': 'Old Name', 'project': 'Old Project', 'professor': 'Old Prof'}
            new_data_dict: {'name': 'New Name', 'project': 'New Project', 'professor': 'New Prof'}
        """
        
        if output_path is None:
            base_name = os.path.splitext(self.doc_path)[0]
            output_path = f"{base_name}_COMPLETE_UPDATE.docx"
        
        print("🚀 Starting comprehensive project document update...")
        print("=" * 70)
        
        # Replace all specified fields
        for field_type, old_value in old_data_dict.items():
            if old_value and new_data_dict.get(field_type):
                new_value = new_data_dict[field_type]
                print(f"\n🔄 Processing {field_type.upper()} replacement:")
                print(f"   OLD: {old_value}")
                print(f"   NEW: {new_value}")
                
                self._replace_text_globally(old_value, new_value, field_type)
        
        # Save the document
        self.doc.save(output_path)
        
        # Print summary
        print("=" * 70)
        print(f"✅ DOCUMENT UPDATE COMPLETE!")
        print(f"📊 Total replacements: {self.replacements}")
        print(f"💾 Saved as: {output_path}")
        
        if self.changes_log:
            print("\n📋 Change Log:")
            for log in self.changes_log[-10:]:  # Show last 10 changes
                print(f"   • {log}")
        
        return self.replacements
    
    def _replace_text_globally(self, old_text, new_text, field_type):
        """Replace text throughout entire document with formatting preservation"""
        
        # Replace in main document paragraphs
        for i, paragraph in enumerate(self.doc.paragraphs):
            if paragraph.text.strip():
                self._smart_replace_with_formatting(paragraph, old_text, new_text, 
                                                  f"paragraph_{i+1}", field_type)
        
        # Replace in tables
        for table_idx, table in enumerate(self.doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        if paragraph.text.strip():
                            location = f"table_{table_idx+1}.cell_{cell_idx+1}.para_{para_idx+1}"
                            self._smart_replace_with_formatting(paragraph, old_text, new_text, 
                                                              location, field_type)
        
        # Replace in headers and footers
        for section_idx, section in enumerate(self.doc.sections):
            # Headers
            if section.header is not None:
                for para_idx, paragraph in enumerate(section.header.paragraphs):
                    if paragraph.text.strip():
                        location = f"header_section_{section_idx+1}.para_{para_idx+1}"
                        self._smart_replace_with_formatting(paragraph, old_text, new_text, 
                                                          location, field_type)
            
            # Footers
            if section.footer is not None:
                for para_idx, paragraph in enumerate(section.footer.paragraphs):
                    if paragraph.text.strip():
                        location = f"footer_section_{section_idx+1}.para_{para_idx+1}"
                        self._smart_replace_with_formatting(paragraph, old_text, new_text, 
                                                          location, field_type)
    
    def _smart_replace_with_formatting(self, paragraph, old_text, new_text, location, field_type):
        """Smart replacement preserving original formatting"""
        
        original_text = paragraph.text
        if not original_text or old_text.lower() not in original_text.lower():
            return
        
        # Store original runs and their formatting
        original_runs = []
        for run in paragraph.runs:
            original_runs.append({
                'text': run.text,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'color': run.font.color.rgb if run.font.color else None,
                'all_caps': run.font.all_caps
            })
        
        # Clear the paragraph
        paragraph.clear()
        
        # Use regex to split while keeping the matches
        pattern = re.compile(f'({re.escape(old_text)})', re.IGNORECASE)
        parts = pattern.split(original_text)
        
        current_pos = 0
        
        for part in parts:
            if not part:
                continue
                
            if part.lower() == old_text.lower():
                # This is the text to replace - use original formatting
                formatting = self._get_formatting_at_position(original_runs, current_pos)
                run = paragraph.add_run(new_text)
                self._apply_formatting(run, formatting)
                
                self.replacements += 1
                log_entry = f"{field_type.upper()} in {location}: '{part}' → '{new_text}'"
                self.changes_log.append(log_entry)
                print(f"✓ {log_entry}")
                
            else:
                # Keep original text with original formatting
                self._add_text_with_original_formatting(paragraph, part, original_runs, current_pos)
            
            current_pos += len(part)
    
    def _get_formatting_at_position(self, runs_data, position):
        """Get formatting at specific text position"""
        current_pos = 0
        for run_data in runs_data:
            run_length = len(run_data['text'])
            if position < current_pos + run_length:
                return {
                    'font_name': run_data['font_name'],
                    'font_size': run_data['font_size'],
                    'bold': run_data['bold'],
                    'italic': run_data['italic'],
                    'underline': run_data['underline'],
                    'color': run_data['color'],
                    'all_caps': run_data['all_caps']
                }
            current_pos += run_length
        return None
    
    def _add_text_with_original_formatting(self, paragraph, text, runs_data, start_pos):
        """Add text with original character-by-character formatting"""
        for i, char in enumerate(text):
            pos = start_pos + i
            formatting = self._get_formatting_at_position(runs_data, pos)
            run = paragraph.add_run(char)
            self._apply_formatting(run, formatting)
    
    def _apply_formatting(self, run, formatting):
        """Apply formatting properties to a run"""
        if formatting:
            if formatting['font_name']:
                run.font.name = formatting['font_name']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), formatting['font_name'])
            
            if formatting['font_size']:
                run.font.size = formatting['font_size']
            
            if formatting['bold'] is not None:
                run.bold = formatting['bold']
            
            if formatting['italic'] is not None:
                run.italic = formatting['italic']
            
            if formatting['underline'] is not None:
                run.underline = formatting['underline']
            
            if formatting['color']:
                run.font.color.rgb = formatting['color']
            
            if formatting['all_caps'] is not None:
                run.font.all_caps = formatting['all_caps']


class SmartProjectBookEditor:
    def __init__(self, doc_path):
        self.doc = Document(doc_path)
        self.doc_path = doc_path
        self.replacements = 0
    
    def auto_detect_and_replace(self, new_student_name, new_project_name, new_professor_name=None, new_guide_name=None):
        """Auto-detect project details and replace them"""
        
        print("🔍 Auto-detecting project details...")
        
        # Extract current project details
        current_details = self._extract_project_details()
        
        print("📋 Detected Details:")
        for key, value in current_details.items():
            if value:
                print(f"   {key.upper()}: {value}")
        
        # Prepare replacement dictionary
        updates = {
            'old_name': current_details.get('student_name'),
            'new_name': new_student_name,
            'old_project': current_details.get('project_title'),
            'new_project': new_project_name,
            'old_professor': current_details.get('professor_name'),
            'new_professor': new_professor_name or current_details.get('professor_name'),
            'old_guide': current_details.get('guide_name'),
            'new_guide': new_guide_name or current_details.get('guide_name')
        }
        
        # Remove None values
        updates = {k: v for k, v in updates.items() if v is not None}
        
        return self.complete_update(updates)
    
    def _extract_project_details(self):
        """Extract project details from document content"""
        details = {
            'student_name': None,
            'project_title': None,
            'professor_name': None,
            'guide_name': None,
            'college_name': None
        }
        
        full_text = "\n".join([para.text for para in self.doc.paragraphs if para.text.strip()])
        
        # Improved patterns for extraction
        patterns = {
            'student_name': r'Submitted by\s*\n*(Mr\.?|Ms\.?)?\s*([^\n]+)',
            'project_title': r'[“"]([^"”]+)[”"]|on\s*\n*"([^"]+)"|on\s*\n*([^\n]+?)(?=\n\s*[A-Z]|$)',
            'professor_name': r'under the guidance of\s*\n*(Prof\.?|Dr\.?)?\s*([^\n]+)',
            'guide_name': r'guide[d]? by\s*\n*(Mr\.?|Ms\.?|Prof\.?|Dr\.?)?\s*([^\n]+)',
            'college_name': r'([A-Z][A-Za-z\s]+COLLEGE|[A-Z][A-Za-z\s]+UNIVERSITY)'
        }
        
        for key, pattern in patterns.items():
            match = re.search(pattern, full_text, re.IGNORECASE | re.MULTILINE)
            if match:
                # Extract the actual value from match groups
                groups = match.groups()
                if groups:
                    # Find the first non-empty group that contains meaningful text
                    for group in groups:
                        if group and len(group.strip()) > 2:
                            details[key] = group.strip()
                            # Special handling for professor/guide names to get full name
                            if key in ['professor_name', 'guide_name'] and len(details[key]) < 4:
                                continue
                            break
        
        # Improved professor name extraction
        if not details['professor_name'] or len(details['professor_name']) < 4:
            prof_match = re.search(r'guidance of[^,\n]*,?\s*([^,\n]+(?:,|$))', full_text, re.IGNORECASE)
            if prof_match and prof_match.group(1):
                details['professor_name'] = prof_match.group(1).strip()
        
        # Clean up the extracted values
        for key in details:
            if details[key]:
                # Remove common prefixes/suffixes
                details[key] = re.sub(r'^(Mr|Ms|Prof|Dr)\.?\s*', '', details[key])
                details[key] = details[key].strip()
        
        return details
    
    def complete_update(self, updates_dict, output_path=None):
        """Complete project document update with manual specifications"""
        
        if output_path is None:
            base_name = os.path.splitext(self.doc_path)[0]
            output_path = f"{base_name}_SMART_UPDATE.docx"
        
        # Create ProjectDocumentEditor instance
        editor = ProjectDocumentEditor(self.doc_path)
        
        # Prepare old and new data dictionaries
        old_data = {}
        new_data = {}
        
        if updates_dict.get('old_name') and updates_dict.get('new_name'):
            old_data['name'] = updates_dict['old_name']
            new_data['name'] = updates_dict['new_name']
        
        if updates_dict.get('old_project') and updates_dict.get('new_project'):
            old_data['project'] = updates_dict['old_project']
            new_data['project'] = updates_dict['new_project']
        
        if updates_dict.get('old_professor') and updates_dict.get('new_professor'):
            old_data['professor'] = updates_dict['old_professor']
            new_data['professor'] = updates_dict['new_professor']
        
        if updates_dict.get('old_guide') and updates_dict.get('new_guide'):
            old_data['guide'] = updates_dict['old_guide']
            new_data['guide'] = updates_dict['new_guide']
        
        return editor.replace_project_details(old_data, new_data, output_path)


# Easy-to-use functions
def quick_project_update(doc_path, new_project_name, output_path=None):
    """Quickly update just the project name"""
    editor = SmartProjectBookEditor(doc_path)
    current_details = editor._extract_project_details()
    
    updates = {
        'old_project': current_details.get('project_title'),
        'new_project': new_project_name
    }
    
    return editor.complete_update(updates, output_path)


def quick_student_update(doc_path, new_student_name, output_path=None):
    """Quickly update just the student name"""
    editor = SmartProjectBookEditor(doc_path)
    current_details = editor._extract_project_details()
    
    updates = {
        'old_name': current_details.get('student_name'),
        'new_name': new_student_name
    }
    
    return editor.complete_update(updates, output_path)


def quick_professor_update(doc_path, new_professor_name, output_path=None):
    """Quickly update just the professor name"""
    editor = SmartProjectBookEditor(doc_path)
    current_details = editor._extract_project_details()
    
    updates = {
        'old_professor': current_details.get('professor_name'),
        'new_professor': new_professor_name
    }
    
    return editor.complete_update(updates, output_path)


def quick_guide_update(doc_path, new_guide_name, output_path=None):
    """Quickly update just the guide name"""
    editor = SmartProjectBookEditor(doc_path)
    current_details = editor._extract_project_details()
    
    updates = {
        'old_guide': current_details.get('guide_name'),
        'new_guide': new_guide_name
    }
    
    return editor.complete_update(updates, output_path)


def complete_project_update(doc_path, updates_dict, output_path=None):
    """
    Complete project document update
    
    Example updates_dict:
    {
        'old_name': 'Param Sachin Shinde',
        'new_name': 'Ankit Rajesh Yadav',
        'old_project': 'The Daily Chronicle', 
        'new_project': 'AI Document Generator',
        'old_professor': 'Prof. Srimathi Narayanan',
        'new_professor': 'Dr. Smith Johnson',
        'old_guide': 'Swetha Iyer',
        'new_guide': 'Prerna Patil'
    }
    """
    editor = SmartProjectBookEditor(doc_path)
    return editor.complete_update(updates_dict, output_path)


# Usage examples
if __name__ == "__main__":
    # Method 1: Auto-detect and replace (FIXED)
    print("=== METHOD 1: AUTO-DETECT AND REPLACE ===")
    editor = SmartProjectBookEditor("proj.docx")
    editor.auto_detect_and_replace(
        new_student_name="Ankit Rajesh Yadav",
        new_project_name="AI-Powered Project Book Generator",
        new_professor_name="Dr. Smith Johnson",
        new_guide_name="Prerna Patil"
    )
    
    print("\n" + "="*70 + "\n")
    
    # Method 2: Quick project name update only
    print("=== METHOD 2: PROJECT NAME UPDATE ===")
    quick_project_update(
        "proj.docx",
        "Advanced Document Processing System"
    )
    
    print("\n" + "="*70 + "\n")
    
    # Method 3: Quick student name update only
    print("=== METHOD 3: STUDENT NAME UPDATE ===")
    quick_student_update(
        "proj.docx", 
        "Param Sachin Shinde"
    )
    
    print("\n" + "="*70 + "\n")
    
    # Method 4: Complete manual update
    print("=== METHOD 4: COMPLETE MANUAL UPDATE ===")
    updates = {
        'old_name': 'Param Sachin Shinde',
        'new_name': 'Rahul Kumar Singh',
        'old_project': 'Face Recognition Attendance',
        'new_project': 'Smart Attendance System with AI',
        'old_professor': 'Prof. Srimathi Narayanan',
        'new_professor': 'Dr. Rajesh Mehta',
        'old_guide': 'Swetha Iyer',
        'new_guide': 'Prerna Patil'
    }
    complete_project_update("proj.docx", updates)