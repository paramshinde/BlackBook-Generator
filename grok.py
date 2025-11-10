from docx import Document
import re
import os
from docx.oxml.ns import qn

class ProjectDocumentEditor:
    def __init__(self, doc_path):
        self.doc_path = doc_path
        self.doc = Document(doc_path)
        self.replacements = 0
        self.changes_log = []
 
    def replace_project_details(self, old_data_dict, new_data_dict, output_path=None):
        """
        Replace multiple project details throughout the document
        
        Args:
            old_data_dict: {'name': 'Old Name', 'project': 'Old Project', 'professor': 'Old Prof'}
            new_data_dict: {'name': 'New Name', 'project': 'New Project', 'professor': 'New Prof'}
        """
        
        if output_path is None:
            base_name = os.path.splitext(self.doc_path)[0]
            output_path = f"{base_name}_UPDATED.docx"
        
        print("🚀 Starting template document update...")
        print("=" * 70)
        
        # Replace all specified fields
        for field_type, old_value in old_data_dict.items():
            if old_value and new_data_dict.get(field_type):
                new_value = new_data_dict[field_type]
                print(f"\n🔄 Processing {field_type.upper()} replacement:")
                print(f"   OLD: {old_value}")
                print(f"   NEW: {new_value}")
                
                self._replace_text_globally(old_value, new_value, field_type)
        
        # Save the document
        self.doc.save(output_path)
        
        # Print summary
        print("=" * 70)
        print(f"✅ DOCUMENT UPDATE COMPLETE!")
        print(f"📊 Total replacements: {self.replacements}")
        print(f"💾 Saved as: {output_path}")
        
        if self.changes_log:
            print("\n📋 Change Log:")
            for log in self.changes_log[-10:]:  # Show last 10 changes
                print(f"   • {log}")
        
        return self.replacements
 
    def _replace_text_globally(self, old_text, new_text, field_type):
        """Replace text throughout entire document with formatting preservation"""
        
        # Replace in main document paragraphs
        for i, paragraph in enumerate(self.doc.paragraphs):
            if paragraph.text.strip():
                self._smart_replace_with_formatting(paragraph, old_text, new_text, 
                                                    f"paragraph_{i+1}", field_type)
        
        # Replace in tables
        for table_idx, table in enumerate(self.doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        if paragraph.text.strip():
                            location = f"table_{table_idx+1}.cell_{cell_idx+1}.para_{para_idx+1}"
                            self._smart_replace_with_formatting(paragraph, old_text, new_text, 
                                                                location, field_type)
        
        # Replace in headers and footers
        for section_idx, section in enumerate(self.doc.sections):
            # Headers
            if section.header is not None:
                for para_idx, paragraph in enumerate(section.header.paragraphs):
                    if paragraph.text.strip():
                        location = f"header_section_{section_idx+1}.para_{para_idx+1}"
                        self._smart_replace_with_formatting(paragraph, old_text, new_text, 
                                                            location, field_type)
            
            # Footers
            if section.footer is not None:
                for para_idx, paragraph in enumerate(section.footer.paragraphs):
                    if paragraph.text.strip():
                        location = f"footer_section_{section_idx+1}.para_{para_idx+1}"
                        self._smart_replace_with_formatting(paragraph, old_text, new_text, 
                                                            location, field_type)
 
    def _smart_replace_with_formatting(self, paragraph, old_text, new_text, location, field_type):
        """Smart replacement preserving original formatting"""
        
        original_text = paragraph.text
        if not original_text or old_text.lower() not in original_text.lower():
            return
        
        # Store original runs and their formatting
        original_runs = []
        for run in paragraph.runs:
            original_runs.append({
                'text': run.text,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'color': run.font.color.rgb if run.font.color else None,
                'all_caps': run.font.all_caps
            })
        
        # Clear the paragraph
        paragraph.clear()
        
        # Use regex to split while keeping the matches
        pattern = re.compile(f'({re.escape(old_text)})', re.IGNORECASE)
        parts = pattern.split(original_text)
        
        current_pos = 0
        
        for part in parts:
            if not part:
                continue
                
            if part.lower() == old_text.lower():
                # This is the text to replace - use original formatting
                formatting = self._get_formatting_at_position(original_runs, current_pos)
                run = paragraph.add_run(new_text)
                self._apply_formatting(run, formatting)
                
                self.replacements += 1
                log_entry = f"{field_type.upper()} in {location}: '{part}' → '{new_text}'"
                self.changes_log.append(log_entry)
                print(f"✓ {log_entry}")
                
            else:
                # Keep original text with original formatting
                self._add_text_with_original_formatting(paragraph, part, original_runs, current_pos)
            
            current_pos += len(part)
 
    def _get_formatting_at_position(self, runs_data, position):
        """Get formatting at specific text position"""
        current_pos = 0
        for run_data in runs_data:
            run_length = len(run_data['text'])
            if position < current_pos + run_length:
                return {
                    'font_name': run_data['font_name'],
                    'font_size': run_data['font_size'],
                    'bold': run_data['bold'],
                    'italic': run_data['italic'],
                    'underline': run_data['underline'],
                    'color': run_data['color'],
                    'all_caps': run_data['all_caps']
                }
            current_pos += run_length
        return None
 
    def _add_text_with_original_formatting(self, paragraph, text, runs_data, start_pos):
        """Add text with original character-by-character formatting"""
        for i, char in enumerate(text):
            pos = start_pos + i
            formatting = self._get_formatting_at_position(runs_data, pos)
            run = paragraph.add_run(char)
            self._apply_formatting(run, formatting)
 
    def _apply_formatting(self, run, formatting):
        """Apply formatting properties to a run"""
        if formatting:
            if formatting['font_name']:
                run.font.name = formatting['font_name']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), formatting['font_name'])
            
            if formatting['font_size']:
                run.font.size = formatting['font_size']
            
            if formatting['bold'] is not None:
                run.bold = formatting['bold']
            
            if formatting['italic'] is not None:
                run.italic = formatting['italic']
            
            if formatting['underline'] is not None:
                run.underline = formatting['underline']
            
            if formatting['color']:
                run.font.color.rgb = formatting['color']
            
            if formatting['all_caps'] is not None:
                run.font.all_caps = formatting['all_caps']


class TemplateProjectEditor:
    def __init__(self, doc_path):
        self.doc_path = doc_path
        self.replacements = 0
 
    def update_template(self, updates_dict, output_path=None):
        """Update the template document with user-provided replacements"""
        
        if output_path is None:
            base_name = os.path.splitext(self.doc_path)[0]
            output_path = f"{base_name}_UPDATED.docx"
        
        # Create ProjectDocumentEditor instance
        editor = ProjectDocumentEditor(self.doc_path)
        
        # Prepare old and new data dictionaries
        old_data = {}
        new_data = {}
        
        if updates_dict.get('new_name'):
            old_data['name'] = 'Param Sachin Shinde'
            new_data['name'] = updates_dict['new_name']
        
        if updates_dict.get('new_project'):
            old_data['project'] = 'Face Recognition Attendance'
            new_data['project'] = updates_dict['new_project']
        
        if updates_dict.get('new_professor'):
            old_data['professor'] = 'Srimathi Narayanan'
            new_data['professor'] = updates_dict['new_professor']
        
        if updates_dict.get('new_guide'):
            old_data['guide'] = 'Swetha Iyer'
            new_data['guide'] = updates_dict['new_guide']
        
        if updates_dict.get('new_college'):
            old_data['college'] = 'vartak college'
            new_data['college'] = updates_dict['new_college']
        
        if updates_dict.get('new_year'):
            old_data['year'] = '2024-2025'
            new_data['year'] = updates_dict['new_year']
        
        return editor.replace_project_details(old_data, new_data, output_path)


# Easy-to-use functions
def quick_project_update(doc_path, new_project_name, output_path=None):
    """Quickly update just the project name"""
    editor = TemplateProjectEditor(doc_path)
    updates = {
        'new_project': new_project_name
    }
    return editor.update_template(updates, output_path)


def quick_student_update(doc_path, new_student_name, output_path=None):
    """Quickly update just the student name"""
    editor = TemplateProjectEditor(doc_path)
    updates = {
        'new_name': new_student_name
    }
    return editor.update_template(updates, output_path)


def quick_professor_update(doc_path, new_professor_name, output_path=None):
    """Quickly update just the professor name"""
    editor = TemplateProjectEditor(doc_path)
    updates = {
        'new_professor': new_professor_name
    }
    return editor.update_template(updates, output_path)


def quick_guide_update(doc_path, new_guide_name, output_path=None):
    """Quickly update just the guide name"""
    editor = TemplateProjectEditor(doc_path)
    updates = {
        'new_guide': new_guide_name
    }
    return editor.update_template(updates, output_path)


def quick_college_update(doc_path, new_college_name, output_path=None):
    """Quickly update just the college name"""
    editor = TemplateProjectEditor(doc_path)
    updates = {
        'new_college': new_college_name
    }
    return editor.update_template(updates, output_path)


def quick_year_update(doc_path, new_year, output_path=None):
    """Quickly update just the year"""
    editor = TemplateProjectEditor(doc_path)
    updates = {
        'new_year': new_year
    }
    return editor.update_template(updates, output_path)


def complete_template_update(doc_path, updates_dict, output_path=None):
    """
    Complete template document update
    
    Example updates_dict:
    {
        'new_name': 'Ankit Rajesh Yadav',
        'new_project': 'AI Document Generator',
        'new_professor': 'Dr. Smith Johnson',
        'new_guide': 'Prerna Patil',
        'new_college': 'New College',
        'new_year': '2025-2026'
    }
    """
    editor = TemplateProjectEditor(doc_path)
    return editor.update_template(updates_dict, output_path)


# Interactive User Input Functions
def interactive_template_update(doc_path):
    """Interactive interface for template updates"""
    print("🎯 Project Template Update Tool")
    print("=" * 50)
    
    # Check if document exists
    if not os.path.exists(doc_path):
        print(f"❌ Error: Document '{doc_path}' not found!")
        return
    
    print(f"\n📄 Using template: {doc_path}")
    print("\n🔍 Enter the new values to replace the hardcoded old values.")
    print("   Press Enter to skip any field.")
    print("\nHardcoded Old Values:")
    print("   Student Name: Param Sachin Shinde")
    print("   Project Title: Face Recognition Attendance")
    print("   Professor Name: Srimathi Narayanan")
    print("   Guide Name: Swetha Iyer")
    print("   College Name: vartak college")
    print("   Year: 2024-2025")
    
    print("\n" + "=" * 50)
    print("Choose update method:")
    print("1. Quick Student Name Update")
    print("2. Quick Project Title Update")
    print("3. Quick Professor Name Update")
    print("4. Quick Guide Name Update")
    print("5. Quick College Name Update")
    print("6. Quick Year Update")
    print("7. Complete Custom Update (All fields)")
    print("8. Exit")
    
    while True:
        try:
            choice = input("\nEnter your choice (1-8): ").strip()
            
            if choice == '1':
                return quick_interactive_student_update(doc_path)
            elif choice == '2':
                return quick_interactive_project_update(doc_path)
            elif choice == '3':
                return quick_interactive_professor_update(doc_path)
            elif choice == '4':
                return quick_interactive_guide_update(doc_path)
            elif choice == '5':
                return quick_interactive_college_update(doc_path)
            elif choice == '6':
                return quick_interactive_year_update(doc_path)
            elif choice == '7':
                return interactive_complete_update(doc_path)
            elif choice == '8':
                print("👋 Goodbye!")
                return
            else:
                print("❌ Invalid choice! Please enter 1-8.")
                
        except KeyboardInterrupt:
            print("\n\n👋 Operation cancelled by user.")
            return


def quick_interactive_student_update(doc_path):
    """Interactive student name update"""
    print("\n👤 STUDENT NAME UPDATE")
    print("-" * 30)
    print("Old Student Name: Param Sachin Shinde")
    
    new_name = input("Enter new student name (or press Enter to skip): ").strip()
    if not new_name:
        print("⏭️ Skipping update.")
        return
    
    print(f"\n🔄 Updating 'Param Sachin Shinde' → '{new_name}'")
    updates = {
        'new_name': new_name
    }
    
    return complete_template_update(doc_path, updates)


def quick_interactive_project_update(doc_path):
    """Interactive project title update"""
    print("\n📋 PROJECT TITLE UPDATE")
    print("-" * 30)
    print("Old Project Title: Face Recognition Attendance")
    
    new_project = input("Enter new project title (or press Enter to skip): ").strip()
    if not new_project:
        print("⏭️ Skipping update.")
        return
    
    print(f"\n🔄 Updating 'Face Recognition Attendance' → '{new_project}'")
    updates = {
        'new_project': new_project
    }
    
    return complete_template_update(doc_path, updates)


def quick_interactive_professor_update(doc_path):
    """Interactive professor name update"""
    print("\n👨‍🏫 PROFESSOR NAME UPDATE")
    print("-" * 30)
    print("Old Professor Name: Srimathi Narayanan")
    
    new_prof = input("Enter new professor name (or press Enter to skip): ").strip()
    if not new_prof:
        print("⏭️ Skipping update.")
        return
    
    print(f"\n🔄 Updating 'Srimathi Narayanan' → '{new_prof}'")
    updates = {
        'new_professor': new_prof
    }
    
    return complete_template_update(doc_path, updates)


def quick_interactive_guide_update(doc_path):
    """Interactive guide name update"""
    print("\n👨‍🏫 GUIDE NAME UPDATE")
    print("-" * 30)
    print("Old Guide Name: Swetha Iyer")
    
    new_guide = input("Enter new guide name (or press Enter to skip): ").strip()
    if not new_guide:
        print("⏭️ Skipping update.")
        return
    
    print(f"\n🔄 Updating 'Swetha Iyer' → '{new_guide}'")
    updates = {
        'new_guide': new_guide
    }
    
    return complete_template_update(doc_path, updates)


def quick_interactive_college_update(doc_path):
    """Interactive college name update"""
    print("\n🏫 COLLEGE NAME UPDATE")
    print("-" * 30)
    print("Old College Name: vartak college")
    
    new_college = input("Enter new college name (or press Enter to skip): ").strip()
    if not new_college:
        print("⏭️ Skipping update.")
        return
    
    print(f"\n🔄 Updating 'vartak college' → '{new_college}'")
    updates = {
        'new_college': new_college
    }
    
    return complete_template_update(doc_path, updates)


def quick_interactive_year_update(doc_path):
    """Interactive year update"""
    print("\n📅 YEAR UPDATE")
    print("-" * 30)
    print("Old Year: 2024-2025")
    
    new_year = input("Enter new year (or press Enter to skip): ").strip()
    if not new_year:
        print("⏭️ Skipping update.")
        return
    
    print(f"\n🔄 Updating '2024-2025' → '{new_year}'")
    updates = {
        'new_year': new_year
    }
    
    return complete_template_update(doc_path, updates)


def interactive_complete_update(doc_path):
    """Interactive complete update"""
    print("\n🔄 COMPLETE TEMPLATE UPDATE")
    print("-" * 40)
    
    updates = {}
    
    # Student Name
    print("\n👤 STUDENT NAME:")
    print("Old: Param Sachin Shinde")
    new_name = input("Enter new student name (or press Enter to skip): ").strip()
    if new_name:
        updates['new_name'] = new_name
    
    # Project Title
    print("\n📋 PROJECT TITLE:")
    print("Old: Face Recognition Attendance")
    new_project = input("Enter new project title (or press Enter to skip): ").strip()
    if new_project:
        updates['new_project'] = new_project
    
    # Professor Name
    print("\n👨‍🏫 PROFESSOR NAME:")
    print("Old: Srimathi Narayanan")
    new_prof = input("Enter new professor name (or press Enter to skip): ").strip()
    if new_prof:
        updates['new_professor'] = new_prof
    
    # Guide Name
    print("\n👨‍🏫 GUIDE NAME:")
    print("Old: Swetha Iyer")
    new_guide = input("Enter new guide name (or press Enter to skip): ").strip()
    if new_guide:
        updates['new_guide'] = new_guide
    
    # College Name
    print("\n🏫 COLLEGE NAME:")
    print("Old: vartak college")
    new_college = input("Enter new college name (or press Enter to skip): ").strip()
    if new_college:
        updates['new_college'] = new_college
    
    # Year
    print("\n📅 YEAR:")
    print("Old: 2024-2025")
    new_year = input("Enter new year (or press Enter to skip): ").strip()
    if new_year:
        updates['new_year'] = new_year
    
    if not updates:
        print("❌ No updates specified!")
        return
    
    # Confirm updates
    print("\n📋 SUMMARY OF UPDATES:")
    print("-" * 30)
    if 'new_name' in updates:
        print(f"🔄 Name: 'Param Sachin Shinde' → '{updates['new_name']}'")
    if 'new_project' in updates:
        print(f"🔄 Project: 'Face Recognition Attendance' → '{updates['new_project']}'")
    if 'new_professor' in updates:
        print(f"🔄 Professor: 'Srimathi Narayanan' → '{updates['new_professor']}'")
    if 'new_guide' in updates:
        print(f"🔄 Guide: 'Swetha Iyer' → '{updates['new_guide']}'")
    if 'new_college' in updates:
        print(f"🔄 College: 'vartak college' → '{updates['new_college']}'")
    if 'new_year' in updates:
        print(f"🔄 Year: '2024-2025' → '{updates['new_year']}'")
    
    confirm = input("\nProceed with these changes? (y/n): ").strip().lower()
    if confirm in ['y', 'yes']:
        return complete_template_update(doc_path, updates)
    else:
        print("❌ Update cancelled.")
        return


# Enhanced main section with interactive mode
if __name__ == "__main__":
    print("🎯 Project Template Update Tool v4.0")
    print("=" * 50)
    
    # Check for command line argument
    import sys
    if len(sys.argv) > 1:
        doc_path = sys.argv[1]
        if not os.path.exists(doc_path):
            print(f"❌ Error: Document '{doc_path}' not found!")
            sys.exit(1)
        print(f"📄 Using template: {doc_path}")
        interactive_template_update(doc_path)
    else:
        # Interactive file selection
        print("📁 Please provide the template document path:")
        doc_path = input("Enter document path (or press Enter for 'proj.docx'): ").strip()
        if not doc_path:
            doc_path = "proj.docx"
        
        if os.path.exists(doc_path):
            interactive_template_update(doc_path)
        else:
            print(f"❌ Error: Document '{doc_path}' not found!")
            print("\n💡 Tip: Drag and drop the .docx file path into the terminal.")
    
    print("\n" + "="*70)
    print("✨ Update complete! Check the output file for changes.")

from docx import Document
import re
import os
from docx.oxml.ns import qn

class ProjectDocumentEditor:
    def __init__(self, doc_path):
        self.doc_path = doc_path
        self.doc = Document(doc_path)
        self.replacements = 0
        self.changes_log = []
 
    def replace_project_details(self, old_data_dict, new_data_dict, output_path=None):
        """
        Replace multiple project details throughout the document
        
        Args:
            old_data_dict: {'name': 'Old Name', 'project': 'Old Project', 'professor': 'Old Prof'}
            new_data_dict: {'name': 'New Name', 'project': 'New Project', 'professor': 'New Prof'}
        """
        
        if output_path is None:
            base_name = os.path.splitext(self.doc_path)[0]
            output_path = f"{base_name}_UPDATED.docx"
        
        print("🚀 Starting template document update...")
        print("=" * 70)
        
        # Replace all specified fields
        for field_type, old_value in old_data_dict.items():
            if old_value and new_data_dict.get(field_type):
                new_value = new_data_dict[field_type]
                print(f"\n🔄 Processing {field_type.upper()} replacement:")
                print(f"   OLD: {old_value}")
                print(f"   NEW: {new_value}")
                
                self._replace_text_globally(old_value, new_value, field_type)
        
        # Save the document
        self.doc.save(output_path)
        
        # Print summary
        print("=" * 70)
        print(f"✅ DOCUMENT UPDATE COMPLETE!")
        print(f"📊 Total replacements: {self.replacements}")
        print(f"💾 Saved as: {output_path}")
        
        if self.changes_log:
            print("\n📋 Change Log:")
            for log in self.changes_log[-10:]:  # Show last 10 changes
                print(f"   • {log}")
        
        return self.replacements
 
    def _replace_text_globally(self, old_text, new_text, field_type):
        """Replace text throughout entire document with formatting preservation"""
        
        # Replace in main document paragraphs
        for i, paragraph in enumerate(self.doc.paragraphs):
            if paragraph.text.strip():
                self._smart_replace_with_formatting(paragraph, old_text, new_text, 
                                                    f"paragraph_{i+1}", field_type)
        
        # Replace in tables
        for table_idx, table in enumerate(self.doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        if paragraph.text.strip():
                            location = f"table_{table_idx+1}.cell_{cell_idx+1}.para_{para_idx+1}"
                            self._smart_replace_with_formatting(paragraph, old_text, new_text, 
                                                                location, field_type)
        
        # Replace in headers and footers
        for section_idx, section in enumerate(self.doc.sections):
            # Headers
            if section.header is not None:
                for para_idx, paragraph in enumerate(section.header.paragraphs):
                    if paragraph.text.strip():
                        location = f"header_section_{section_idx+1}.para_{para_idx+1}"
                        self._smart_replace_with_formatting(paragraph, old_text, new_text, 
                                                            location, field_type)
            
            # Footers
            if section.footer is not None:
                for para_idx, paragraph in enumerate(section.footer.paragraphs):
                    if paragraph.text.strip():
                        location = f"footer_section_{section_idx+1}.para_{para_idx+1}"
                        self._smart_replace_with_formatting(paragraph, old_text, new_text, 
                                                            location, field_type)
 
    def _smart_replace_with_formatting(self, paragraph, old_text, new_text, location, field_type):
        """Smart replacement preserving original formatting"""
        
        original_text = paragraph.text
        if not original_text or old_text.lower() not in original_text.lower():
            return
        
        # Store original runs and their formatting
        original_runs = []
        for run in paragraph.runs:
            original_runs.append({
                'text': run.text,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'color': run.font.color.rgb if run.font.color else None,
                'all_caps': run.font.all_caps
            })
        
        # Clear the paragraph
        paragraph.clear()
        
        # Use regex to split while keeping the matches
        pattern = re.compile(f'({re.escape(old_text)})', re.IGNORECASE)
        parts = pattern.split(original_text)
        
        current_pos = 0
        
        for part in parts:
            if not part:
                continue
                
            if part.lower() == old_text.lower():
                # This is the text to replace - use original formatting
                formatting = self._get_formatting_at_position(original_runs, current_pos)
                run = paragraph.add_run(new_text)
                self._apply_formatting(run, formatting)
                
                self.replacements += 1
                log_entry = f"{field_type.upper()} in {location}: '{part}' → '{new_text}'"
                self.changes_log.append(log_entry)
                print(f"✓ {log_entry}")
                
            else:
                # Keep original text with original formatting
                self._add_text_with_original_formatting(paragraph, part, original_runs, current_pos)
            
            current_pos += len(part)
 
    def _get_formatting_at_position(self, runs_data, position):
        """Get formatting at specific text position"""
        current_pos = 0
        for run_data in runs_data:
            run_length = len(run_data['text'])
            if position < current_pos + run_length:
                return {
                    'font_name': run_data['font_name'],
                    'font_size': run_data['font_size'],
                    'bold': run_data['bold'],
                    'italic': run_data['italic'],
                    'underline': run_data['underline'],
                    'color': run_data['color'],
                    'all_caps': run_data['all_caps']
                }
            current_pos += run_length
        return None
 
    def _add_text_with_original_formatting(self, paragraph, text, runs_data, start_pos):
        """Add text with original character-by-character formatting"""
        for i, char in enumerate(text):
            pos = start_pos + i
            formatting = self._get_formatting_at_position(runs_data, pos)
            run = paragraph.add_run(char)
            self._apply_formatting(run, formatting)
 
    def _apply_formatting(self, run, formatting):
        """Apply formatting properties to a run"""
        if formatting:
            if formatting['font_name']:
                run.font.name = formatting['font_name']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), formatting['font_name'])
            
            if formatting['font_size']:
                run.font.size = formatting['font_size']
            
            if formatting['bold'] is not None:
                run.bold = formatting['bold']
            
            if formatting['italic'] is not None:
                run.italic = formatting['italic']
            
            if formatting['underline'] is not None:
                run.underline = formatting['underline']
            
            if formatting['color']:
                run.font.color.rgb = formatting['color']
            
            if formatting['all_caps'] is not None:
                run.font.all_caps = formatting['all_caps']


class TemplateProjectEditor:
    def __init__(self, doc_path):
        self.doc_path = doc_path
        self.replacements = 0
        
        # Hardcoded old values from the template
        self.template_values = {
            'student_name': 'Param Sachin Shinde',
            'project_title': 'Face Recognition Attendance',
            'professor_name': 'Srimathi Narayanan',
            'guide_name': 'Swetha Iyer',
            'college_name': 'vartak college',
            'year': '2024-2025'
        }
 
    def update_template(self, new_values_dict, output_path=None):
        """Update the template document with user-provided new values"""
        
        if output_path is None:
            base_name = os.path.splitext(self.doc_path)[0]
            output_path = f"{base_name}_UPDATED.docx"
        
        print("📋 Template Values (will be replaced):")
        for key, old_val in self.template_values.items():
            new_val = new_values_dict.get(key, old_val)
            print(f"   {key.replace('_', ' ').title()}: '{old_val}' → '{new_val}'")
        
        # Create ProjectDocumentEditor instance
        editor = ProjectDocumentEditor(self.doc_path)
        
        # Prepare old and new data dictionaries
        old_data = {}
        new_data = {}
        
        for key, old_val in self.template_values.items():
            new_val = new_values_dict.get(key, old_val)
            if new_val != old_val:  # Only replace if changed
                if key == 'student_name':
                    old_data['name'] = old_val
                    new_data['name'] = new_val
                elif key == 'project_title':
                    old_data['project'] = old_val
                    new_data['project'] = new_val
                elif key == 'professor_name':
                    old_data['professor'] = old_val
                    new_data['professor'] = new_val
                elif key == 'guide_name':
                    old_data['guide'] = old_val
                    new_data['guide'] = new_val
                elif key == 'college_name':
                    old_data['college'] = old_val
                    new_data['college'] = new_val
                elif key == 'year':
                    old_data['year'] = old_val
                    new_data['year'] = new_val
        
        return editor.replace_project_details(old_data, new_data, output_path)


# Interactive User Input Functions
def interactive_template_update(doc_path):
    """Interactive interface for template updates"""
    print("🎯 Project Template Update Tool")
    print("=" * 50)
    
    # Check if document exists
    if not os.path.exists(doc_path):
        print(f"❌ Error: Document '{doc_path}' not found!")
        return
    
    print(f"\n📄 Using template: {doc_path}")
    print("\n✨ Enter new values to replace hardcoded template values.")
    print("   Press Enter to keep the current value.\n")
    
    # Initialize editor
    editor = TemplateProjectEditor(doc_path)
    
    # Get user input for new values
    new_values = {}
    
    print("📝 Current Template Values:")
    for key, old_val in editor.template_values.items():
        display_key = key.replace('_', ' ').title()
        new_val = input(f"{display_key} ('{old_val}'): ").strip()
        if new_val:  # If user enters something, use it
            new_values[key] = new_val
        else:  # Keep original if Enter pressed
            new_values[key] = old_val
    
    # Confirm updates
    print("\n📋 SUMMARY OF UPDATES:")
    print("-" * 30)
    changes_made = False
    for key, old_val in editor.template_values.items():
        new_val = new_values.get(key, old_val)
        if new_val != old_val:
            print(f"🔄 {key.replace('_', ' ').title()}: '{old_val}' → '{new_val}'")
            changes_made = True
        else:
            print(f"⏭️  {key.replace('_', ' ').title()}: No change")
    
    if not changes_made:
        print("\n❌ No changes detected! Exiting.")
        return
    
    confirm = input("\nProceed with these changes? (y/n): ").strip().lower()
    if confirm in ['y', 'yes']:
        return editor.update_template(new_values)
    else:
        print("❌ Update cancelled.")
        return


# Enhanced main section with interactive mode
if __name__ == "__main__":
    print("🎯 Project Template Update Tool v4.0")
    print("=" * 50)
    
    # Check for command line argument
    import sys
    if len(sys.argv) > 1:
        doc_path = sys.argv[1]
        if not os.path.exists(doc_path):
            print(f"❌ Error: Document '{doc_path}' not found!")
            sys.exit(1)
        print(f"📄 Using template: {doc_path}")
        interactive_template_update(doc_path)
    else:
        # Interactive file selection
        print("📁 Please provide the template document path:")
        doc_path = input("Enter document path (or press Enter for 'proj.docx'): ").strip()
        if not doc_path:
            doc_path = "proj.docx"
        
        if os.path.exists(doc_path):
            interactive_template_update(doc_path)
        else:
            print(f"❌ Error: Document '{doc_path}' not found!")
            print("\n💡 Tip: Drag and drop the .docx file path into the terminal.")
    
    print("\n" + "="*70)
    print("✨ Update complete! Check the output file for changes.")