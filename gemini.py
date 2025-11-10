from docx import Document
import re
import os
from docx.oxml.ns import qn
import datetime

class ProjectDocumentEditor:
    def __init__(self, doc_path):
        self.doc_path = doc_path
        self.doc = Document(doc_path)
        self.replacements = 0
        self.changes_log = []
    
    def replace_project_details(self, old_data_dict, new_data_dict, output_path=None):
        """
        Replace multiple project details throughout the document
        
        Args:
            old_data_dict: {'name': 'Old Name', 'project': 'Old Project', 'professor': 'Old Prof', 'year': 'Old Year'}
            new_data_dict: {'name': 'New Name', 'project': 'New Project', 'professor': 'New Prof', 'year': 'New Year'}
        """
        
        if output_path is None:
            base_name = os.path.splitext(self.doc_path)[0]
            output_path = f"{base_name}_COMPLETE_UPDATE.docx"
        
        print("🚀 Starting comprehensive project document update...")
        print("=" * 70)
        
        # Replace all specified fields
        for field_type, old_value in old_data_dict.items():
            if old_value and new_data_dict.get(field_type):
                new_value = new_data_dict[field_type]
                print(f"\n🔄 Processing {field_type.upper()} replacement:")
                print(f"   OLD: {old_value}")
                print(f"   NEW: {new_value}")
                
                self._replace_text_globally(old_value, new_value, field_type)
        
        # Save the document
        self.doc.save(output_path)
        
        # Print summary
        print("=" * 70)
        print(f"✅ DOCUMENT UPDATE COMPLETE!")
        print(f"📊 Total replacements: {self.replacements}")
        print(f"💾 Saved as: {output_path}")
        
        if self.changes_log:
            print("\n📋 Change Log:")
            for log in self.changes_log[-10:]:  # Show last 10 changes
                print(f"   • {log}")
        
        return self.replacements
    
    def _replace_text_globally(self, old_text, new_text, field_type):
        """Replace text throughout entire document with formatting preservation"""
        
        # Replace in main document paragraphs
        for i, paragraph in enumerate(self.doc.paragraphs):
            if paragraph.text.strip():
                self._smart_replace_with_formatting(paragraph, old_text, new_text, 
                                                    f"paragraph_{i+1}", field_type)
        
        # Replace in tables
        for table_idx, table in enumerate(self.doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        if paragraph.text.strip():
                            location = f"table_{table_idx+1}.cell_{cell_idx+1}.para_{para_idx+1}"
                            self._smart_replace_with_formatting(paragraph, old_text, new_text, 
                                                                location, field_type)
        
        # Replace in headers and footers
        for section_idx, section in enumerate(self.doc.sections):
            # Headers
            if section.header is not None:
                for para_idx, paragraph in enumerate(section.header.paragraphs):
                    if paragraph.text.strip():
                        location = f"header_section_{section_idx+1}.para_{para_idx+1}"
                        self._smart_replace_with_formatting(paragraph, old_text, new_text, 
                                                            location, field_type)
            
            # Footers
            if section.footer is not None:
                for para_idx, paragraph in enumerate(section.footer.paragraphs):
                    if paragraph.text.strip():
                        location = f"footer_section_{section_idx+1}.para_{para_idx+1}"
                        self._smart_replace_with_formatting(paragraph, old_text, new_text, 
                                                            location, field_type)
    
    def _smart_replace_with_formatting(self, paragraph, old_text, new_text, location, field_type):
        """Smart replacement preserving original formatting"""
        
        original_text = paragraph.text
        if not original_text or old_text.lower() not in original_text.lower():
            return
        
        # Store original runs and their formatting
        original_runs = []
        for run in paragraph.runs:
            original_runs.append({
                'text': run.text,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'color': run.font.color.rgb if run.font.color else None,
                'all_caps': run.font.all_caps
            })
        
        # Clear the paragraph
        paragraph.clear()
        
        # Use regex to split while keeping the matches
        pattern = re.compile(f'({re.escape(old_text)})', re.IGNORECASE)
        parts = pattern.split(original_text)
        
        current_pos = 0
        
        for part in parts:
            if not part:
                continue
                
            if part.lower() == old_text.lower():
                # This is the text to replace - use original formatting
                formatting = self._get_formatting_at_position(original_runs, current_pos)
                run = paragraph.add_run(new_text)
                self._apply_formatting(run, formatting)
                
                self.replacements += 1
                log_entry = f"{field_type.upper()} in {location}: '{part}' → '{new_text}'"
                self.changes_log.append(log_entry)
                print(f"✓ {log_entry}")
                
            else:
                # Keep original text with original formatting
                self._add_text_with_original_formatting(paragraph, part, original_runs, current_pos)
            
            current_pos += len(part)
    
    def _get_formatting_at_position(self, runs_data, position):
        """Get formatting at specific text position"""
        current_pos = 0
        for run_data in runs_data:
            run_length = len(run_data['text'])
            if position < current_pos + run_length:
                return {
                    'font_name': run_data['font_name'],
                    'font_size': run_data['font_size'],
                    'bold': run_data['bold'],
                    'italic': run_data['italic'],
                    'underline': run_data['underline'],
                    'color': run_data['color'],
                    'all_caps': run_data['all_caps']
                }
            current_pos += run_length
        return None
    
    def _add_text_with_original_formatting(self, paragraph, text, runs_data, start_pos):
        """Add text with original character-by-character formatting"""
        for i, char in enumerate(text):
            pos = start_pos + i
            formatting = self._get_formatting_at_position(runs_data, pos)
            run = paragraph.add_run(char)
            self._apply_formatting(run, formatting)
    
    def _apply_formatting(self, run, formatting):
        """Apply formatting properties to a run"""
        if formatting:
            if formatting['font_name']:
                run.font.name = formatting['font_name']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), formatting['font_name'])
            
            if formatting['font_size']:
                run.font.size = formatting['font_size']
            
            if formatting['bold'] is not None:
                run.bold = formatting['bold']
            
            if formatting['italic'] is not None:
                run.italic = formatting['italic']
            
            if formatting['underline'] is not None:
                run.underline = formatting['underline']
            
            if formatting['color']:
                run.font.color.rgb = formatting['color']
            
            if formatting['all_caps'] is not None:
                run.font.all_caps = formatting['all_caps']


class SmartProjectBookEditor:
    def __init__(self, doc_path):
        self.doc = Document(doc_path)
        self.doc_path = doc_path
        self.replacements = 0
    
    def complete_update(self, updates_dict, output_path=None):
        """Complete project document update with manual specifications"""
        
        if output_path is None:
            base_name = os.path.splitext(self.doc_path)[0]
            output_path = f"{base_name}_COMPLETE_UPDATE.docx"
        
        editor = ProjectDocumentEditor(self.doc_path)
        
        old_data = {}
        new_data = {}
        
        if updates_dict.get('old_name') and updates_dict.get('new_name'):
            old_data['name'] = updates_dict['old_name']
            new_data['name'] = updates_dict['new_name']
        
        if updates_dict.get('old_project') and updates_dict.get('new_project'):
            old_data['project'] = updates_dict['old_project']
            new_data['project'] = updates_dict['new_project']
        
        if updates_dict.get('old_professor') and updates_dict.get('new_professor'):
            old_data['professor'] = updates_dict['old_professor']
            new_data['professor'] = updates_dict['new_professor']
        
        if updates_dict.get('old_guide') and updates_dict.get('new_guide'):
            old_data['guide'] = updates_dict['old_guide']
            new_data['guide'] = updates_dict['new_guide']
        
        if updates_dict.get('old_college') and updates_dict.get('new_college'):
            old_data['college'] = updates_dict['old_college']
            new_data['college'] = updates_dict['new_college']
            
        if updates_dict.get('old_year') and updates_dict.get('new_year'):
            old_data['year'] = updates_dict['old_year']
            new_data['year'] = updates_dict['new_year']
        
        return editor.replace_project_details(old_data, new_data, output_path)

# Main script to run the update process
def main():
    doc_path = os.path.join(os.getcwd(), 'ankit.docx')

    if not os.path.exists(doc_path):
        print(f"Error: The template file '{doc_path}' was not found.")
        print("Please ensure 'ankit.docx' is in the same directory as this script.")
        return

    # Create the editor object here, outside the conditional block
    editor = SmartProjectBookEditor(doc_path)

    # Hardcoded old details from the template
    hardcoded_old_details = {
        'student_name': 'Ankit Rajesh Yadav',
        'project_title': 'The Daily Chronicle',
        'professor_name': 'Srimathi Narayanan',
        'guide_name': 'Srimathi Narayanan',
        'college_name': 'vartak college',
        'year': '2024-2025'
    }
    
    # Generate the default new year
    current_year = datetime.datetime.now().year
    next_year = current_year + 1
    default_new_year = f"{current_year}-{next_year}"
    
    print("\n📝 Template Details Detected from 'ankit.docx':")
    for key, value in hardcoded_old_details.items():
        print(f"  {key.replace('_', ' ').title()}: {value}")
    
    print("\n✨ Please enter the new details. Press Enter to keep the old value.")
    
    new_student_name = input(f"Enter new Student Name (current: {hardcoded_old_details['student_name']}): ") or hardcoded_old_details['student_name']
    new_project_name = input(f"Enter new Project Title (current: {hardcoded_old_details['project_title']}): ") or hardcoded_old_details['project_title']
    new_professor_name = input(f"Enter new Professor Name (current: {hardcoded_old_details['professor_name']}): ") or hardcoded_old_details['professor_name']
    new_guide_name = input(f"Enter new Guide Name (current: {hardcoded_old_details['guide_name']}): ") or hardcoded_old_details['guide_name']
    new_college_name = input(f"Enter new College Name (current: {hardcoded_old_details['college_name']}): ") or hardcoded_old_details['college_name']
    new_year = input(f"Enter new Year (e.g., {default_new_year}) (current: {hardcoded_old_details['year']}): ") or hardcoded_old_details['year']

    updates = {
        'old_name': hardcoded_old_details['student_name'],
        'new_name': new_student_name,
        'old_project': hardcoded_old_details['project_title'],
        'new_project': new_project_name,
        'old_professor': hardcoded_old_details['professor_name'],
        'new_professor': new_professor_name,
        'old_guide': hardcoded_old_details['guide_name'],
        'new_guide': new_guide_name,
        'old_college': hardcoded_old_details['college_name'],
        'new_college': new_college_name,
        'old_year': hardcoded_old_details['year'],
        'new_year': new_year
    }
    
    final_updates = {}
    for key, value in updates.items():
        if key.startswith('old_') and value is not None:
            new_key = 'new_' + key[4:]
            if updates.get(new_key) and updates.get(new_key) != value:
                final_updates[key] = value
                final_updates[new_key] = updates[new_key]

    if final_updates:
        editor.complete_update(final_updates)
    else:
        print("\n🤔 No changes were entered. Document was not modified.")
        

if __name__ == "__main__":
    main()